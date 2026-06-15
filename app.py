import html
import csv
import hashlib
import io
import re
import time
import zipfile
from collections import Counter
from dataclasses import dataclass
from datetime import date
from typing import Any, Dict, List, Optional, Set, Tuple
from xml.etree import ElementTree as ET

ROWS = 30
COLS = 28
TOTAL = ROWS * COLS


@dataclass
class ParseResult:
    floor_ids: List[int]
    glove_names: List[str]
    declared_floor_count: Optional[int] = None
    floor_count_mismatch: bool = False


def _safe_int(text: str) -> Optional[int]:
    try:
        return int(text)
    except Exception:
        return None


def _compact_marker_text(text: str) -> str:
    return re.sub(r"[\s:：\u3000]+", "", text)


def _find_marker_span(lines: List[str], markers: List[str], window: int = 8) -> Optional[Tuple[int, int]]:
    if not lines:
        return None
    compact_lines = [_compact_marker_text(line) for line in lines]
    for start in range(len(compact_lines)):
        merged = ""
        end_limit = min(len(compact_lines), start + window)
        for end in range(start, end_limit):
            merged += compact_lines[end]
            for marker in markers:
                if merged.startswith(marker):
                    return (start, end + 1)
    return None


def extract_declared_floor_count(text: str) -> Optional[int]:
    compact = _compact_marker_text(text)
    match = re.search(
        r"(?:辐射超标地面网格个数|辐射超标地面点位个数|辐射超标地面点位数量|辐射超标地面网格数)(\d{1,4})",
        compact,
    )
    if not match:
        return None

    value = _safe_int(match.group(1))
    if value is None or not (0 <= value <= TOTAL):
        return None
    return value


def parse_report_text(text: str) -> ParseResult:
    """Parse floor-point IDs and glove-hole names from report text."""
    normalized = text.replace("\r", "\n")
    candidates: List[int] = []
    lines = [line.strip() for line in normalized.split("\n") if line.strip()]

    floor_span = _find_marker_span(lines, ["辐射超标地面点位ID", "地面点位ID"])
    glove_span = _find_marker_span(lines, ["辐射超标手套孔名称", "手套孔名称"])

    section_lines: List[str] = []
    if floor_span:
        section_start = floor_span[1]
        section_end = len(lines)
        if glove_span and glove_span[0] > section_start:
            section_end = glove_span[0]
        section_lines = lines[section_start:section_end]

    if section_lines:
        section_text = "\n".join(section_lines)
        section_tokens = [token.strip() for token in section_lines if token.strip()]
        is_digits = re.compile(r"^\d+$").match
        is_decimal_split = re.compile(r"^\.\d+$").match
        is_decimal_full = re.compile(r"^\d+\.\d+$").match

        # Table-like rows can be split into many tokens, e.g.:
        # 53 / 6 / 0 / .12  -> ID=536
        # 566 / 0 / .12     -> ID=566
        i = 0
        while i < len(section_tokens):
            if not is_digits(section_tokens[i]):
                i += 1
                continue

            start = i
            end = i
            while end < len(section_tokens) and is_digits(section_tokens[end]):
                end += 1

            id_text: Optional[str] = None
            consume_to = end
            if end < len(section_tokens) and is_decimal_split(section_tokens[end]):
                if end - start >= 2:
                    id_text = "".join(section_tokens[start : end - 1])
                consume_to = end + 1
            elif end < len(section_tokens) and is_decimal_full(section_tokens[end]):
                id_text = "".join(section_tokens[start:end])
                consume_to = end + 1
            elif end + 1 < len(section_tokens) and section_tokens[end] == "." and is_digits(
                section_tokens[end + 1]
            ):
                if end - start >= 2:
                    id_text = "".join(section_tokens[start : end - 1])
                consume_to = end + 2

            if id_text:
                value = _safe_int(id_text)
                if value is not None and 1 <= value <= TOTAL:
                    candidates.append(value)
                i = consume_to
                continue

            i = end

        # List-style reports may provide only IDs without measurement columns.
        if not candidates and "." not in section_text:
            for token in section_lines:
                if re.fullmatch(r"\d{1,4}", token):
                    value = _safe_int(token)
                    if value is not None and 1 <= value <= TOTAL:
                        candidates.append(value)

    if not candidates:
        # Conservative fallback for non-standard text; require explicit "点位ID" prefix.
        for token in re.findall(r"点位\s*ID\s*[:：]?\s*(\d{1,4})", normalized):
            value = _safe_int(token)
            if value is not None and 1 <= value <= TOTAL:
                candidates.append(value)

    glove_names = re.findall(r"手套孔\s*\d+\s*#", normalized)
    if not glove_names:
        glove_section = re.search(r"(?:辐射超标手套孔名称|手套孔名称)[\s\S]*$", normalized)
        if glove_section:
            glove_names = re.findall(r"手套孔\s*\d+\s*#", glove_section.group(0))

    unique_floor_ids = list(dict.fromkeys(candidates))
    declared_floor_count = extract_declared_floor_count(normalized)
    floor_count_mismatch = (
        declared_floor_count is not None and declared_floor_count != len(unique_floor_ids)
    )

    return ParseResult(
        floor_ids=unique_floor_ids,
        glove_names=list(dict.fromkeys(name.replace(" ", "") for name in glove_names)),
        declared_floor_count=declared_floor_count,
        floor_count_mismatch=floor_count_mismatch,
    )


def parse_point_ids(text: str) -> List[int]:
    """Extract valid point IDs (1..840) from arbitrary text."""
    ids: List[int] = []
    for token in re.findall(r"\b\d{1,4}\b", text):
        value = _safe_int(token)
        if value is not None and 1 <= value <= TOTAL:
            ids.append(value)
    return list(dict.fromkeys(ids))


def id_to_row_col(point_id: int) -> Tuple[int, int]:
    return ((point_id - 1) % ROWS + 1, (point_id - 1) // ROWS + 1)


def row_col_to_id(row: int, col: int) -> int:
    return row + (col - 1) * ROWS


def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()

    if lower.endswith(".txt"):
        for enc in ("utf-8", "gbk", "gb18030"):
            try:
                return file_bytes.decode(enc)
            except Exception:
                continue
        return file_bytes.decode("utf-8", errors="ignore")

    if lower.endswith(".docx"):
        def _extract_docx_xml_text() -> str:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                xml_data = zf.read("word/document.xml")
            root = ET.fromstring(xml_data)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            texts = [node.text for node in root.findall(".//w:t", ns) if node.text]
            return "\n".join(texts)

        def _collect_table_text(table: Any) -> List[str]:
            chunks: List[str] = []
            for row in table.rows:
                for cell in row.cells:
                    cell_lines = [p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()]
                    if cell_lines:
                        chunks.append("\n".join(cell_lines))
                    for nested_table in cell.tables:
                        chunks.extend(_collect_table_text(nested_table))
            return chunks

        try:
            from docx import Document  # type: ignore

            doc = Document(io.BytesIO(file_bytes))
            chunks = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                chunks.extend(_collect_table_text(table))

            merged = "\n".join(chunks).strip()
            if merged:
                return merged
            return _extract_docx_xml_text()
        except Exception:
            return _extract_docx_xml_text()

    if lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise ValueError("PDF 解析失败：请先安装 pypdf，或另存为 txt/docx。") from exc

    raise ValueError("仅支持 .txt / .docx / .pdf 文件")


def _normalized_date(year_text: str, month_text: str, day_text: str) -> Optional[str]:
    year = _safe_int(year_text)
    month = _safe_int(month_text)
    day = _safe_int(day_text)
    if year is None or month is None or day is None:
        return None
    try:
        return date(year, month, day).strftime("%Y%m%d")
    except ValueError:
        return None


def _normalize_date_match_text(text: str) -> str:
    # Some DOCX tables split date digits into separate runs/cells, e.g. "2025-0\\n9-0\\n2".
    # Merge whitespace between adjacent digits before regex matching.
    return re.sub(r"(?<=\d)\s+(?=\d)", "", text)


def extract_date_from_text(text: str) -> Optional[str]:
    normalized_text = _normalize_date_match_text(text)
    patterns = [
        r"(?<!\d)(20\d{2})\s*(?:[\-\._/]|年)\s*(\d{1,2})\s*(?:[\-\._/]|月)\s*(\d{1,2})(?:\s*日)?",
        r"(?<!\d)(20\d{2})\s*(\d{2})\s*(\d{2})",
    ]
    for pattern in patterns:
        for year_text, month_text, day_text in re.findall(pattern, normalized_text):
            normalized = _normalized_date(year_text, month_text, day_text)
            if normalized:
                return normalized
    return None


def extract_task_end_date(text: str) -> Optional[str]:
    normalized_text = _normalize_date_match_text(text)
    label = r"(?:任\s*务\s*)?结\s*束\s*时\s*间\s*[:：]?\s*"
    patterns = [
        label + r"(20\d{2})\s*(?:[年\-\._/])\s*(\d{1,2})\s*(?:[月\-\._/])\s*(\d{1,2})(?:\s*日)?",
        label + r"(20\d{2})\s*(\d{2})\s*(\d{2})",
    ]
    for pattern in patterns:
        match = re.search(pattern, normalized_text)
        if not match:
            continue
        year_text, month_text, day_text = match.groups()
        normalized = _normalized_date(year_text, month_text, day_text)
        if normalized:
            return normalized
    return None


def build_file_signature(filename: str, file_bytes: bytes) -> str:
    digest = hashlib.md5(file_bytes).hexdigest()[:12]
    return f"{filename}|{len(file_bytes)}|{digest}"


def build_text_signature(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()[:12]


def merge_selected_ids(existing_ids: Set[int], new_ids: Set[int]) -> Set[int]:
    return set(existing_ids) | set(new_ids)


def apply_pollution_action(action: str, selected_ids: Set[int], polluted_ids: Set[int]) -> Set[int]:
    new_polluted = set(polluted_ids)
    if action == "set_polluted":
        new_polluted |= selected_ids
    elif action == "set_clean":
        new_polluted -= selected_ids
    return new_polluted


def sync_report_pollution_state(report: Dict[str, Any]) -> ParseResult:
    report_text = str(report.get("report_text", ""))
    parsed = parse_report_text(report_text)
    text_signature = build_text_signature(report_text)

    if report.get("pollution_source_signature") != text_signature:
        report["polluted_ids"] = set(parsed.floor_ids)
        report["pollution_source_signature"] = text_signature
        report["selected_ids"] = set()

    if "polluted_ids" not in report:
        report["polluted_ids"] = set(parsed.floor_ids)
    return parsed


def get_report_polluted_ids(report: Dict[str, Any]) -> Set[int]:
    raw_ids = report.get("polluted_ids")
    if isinstance(raw_ids, set):
        return {point_id for point_id in raw_ids if isinstance(point_id, int) and 1 <= point_id <= TOTAL}
    if isinstance(raw_ids, list):
        return {point_id for point_id in raw_ids if isinstance(point_id, int) and 1 <= point_id <= TOTAL}
    parsed = parse_report_text(str(report.get("report_text", "")))
    return set(parsed.floor_ids)


def build_report_switch_options(
    report_keys: List[str], report_store: Dict[str, Dict[str, Any]]
) -> List[Tuple[str, str]]:
    def sort_key(key: str) -> Tuple[int, str, str]:
        report = report_store.get(key, {})
        report_date = report.get("date")
        file_name = str(report.get("file_name", ""))
        if isinstance(report_date, str) and report_date:
            return (0, report_date, file_name.lower())
        return (1, "99999999", file_name.lower())

    ordered_keys = sorted(report_keys, key=sort_key)
    date_labels = [str(report_store.get(key, {}).get("date") or "未识别日期") for key in ordered_keys]
    counts: Dict[str, int] = {}
    for date_label in date_labels:
        counts[date_label] = counts.get(date_label, 0) + 1

    seen: Dict[str, int] = {}
    options: List[Tuple[str, str]] = []
    for key, base in zip(ordered_keys, date_labels):
        seen[base] = seen.get(base, 0) + 1
        if counts[base] > 1:
            label = f"{base} ({seen[base]})"
        else:
            label = base
        options.append((key, label))
    return options


def build_cross_report_summary(
    report_options: List[Tuple[str, str]], report_store: Dict[str, Dict[str, Any]]
) -> List[Dict[str, Any]]:
    def compact_ids(point_ids: Set[int], limit: int = 16) -> str:
        if not point_ids:
            return "-"
        ordered = sorted(point_ids)
        preview = "、".join(str(point_id) for point_id in ordered[:limit])
        if len(ordered) > limit:
            preview += " ..."
        return preview

    rows: List[Dict[str, Any]] = []
    prev_polluted_ids: Optional[Set[int]] = None

    for report_key, _ in report_options:
        report = report_store.get(report_key, {})
        parsed = parse_report_text(str(report.get("report_text", "")))
        current_polluted_ids = get_report_polluted_ids(report)

        if prev_polluted_ids is None:
            diff_count = "基线"
            added_ids = current_polluted_ids
            removed_ids: Set[int] = set()
        else:
            diff_count = f"{len(current_polluted_ids) - len(prev_polluted_ids):+d}"
            added_ids = current_polluted_ids - prev_polluted_ids
            removed_ids = prev_polluted_ids - current_polluted_ids

        rows.append(
            {
                "日期": str(report.get("date") or "未识别日期"),
                "文件名": str(report.get("file_name", "-")),
                "污染地面点位数": len(current_polluted_ids),
                "超标手套孔数": len(parsed.glove_names),
                "较前一日变化": diff_count,
                "新增点位数": len(added_ids),
                "减少点位数": len(removed_ids),
                "新增点位ID(预览)": compact_ids(added_ids),
                "减少点位ID(预览)": compact_ids(removed_ids),
            }
        )
        prev_polluted_ids = current_polluted_ids

    return rows


def normalize_point_set(raw_ids: Any) -> Set[int]:
    if isinstance(raw_ids, set):
        values = raw_ids
    elif isinstance(raw_ids, list):
        values = set(raw_ids)
    else:
        values = set()
    return {point_id for point_id in values if isinstance(point_id, int) and 1 <= point_id <= TOTAL}


def compute_centroid(point_ids: Set[int]) -> Optional[Tuple[float, float]]:
    if not point_ids:
        return None
    rows = [id_to_row_col(point_id)[0] for point_id in point_ids]
    cols = [id_to_row_col(point_id)[1] for point_id in point_ids]
    return (sum(rows) / len(rows), sum(cols) / len(cols))


def snap_centroid_to_polluted_point(
    point_ids: Set[int], target_row: float, target_col: float
) -> Optional[Tuple[int, int, int]]:
    if not point_ids:
        return None

    best_id = min(
        point_ids,
        key=lambda point_id: (
            (id_to_row_col(point_id)[0] - target_row) ** 2 + (id_to_row_col(point_id)[1] - target_col) ** 2,
            point_id,
        ),
    )
    row, col = id_to_row_col(best_id)
    return (best_id, row, col)


def recommend_centroid_smoothing_window(total_reports: int, max_window: Optional[int] = None) -> int:
    if total_reports <= 0:
        return 1

    limit = max(1, min(int(max_window or total_reports), total_reports))
    recommended = max(1, (total_reports + 4) // 5)
    return min(recommended, limit)


def build_centroid_window_points(history: List[Dict[str, Any]], window_size: int = 1) -> List[Dict[str, Any]]:
    if not history:
        return []

    step = max(1, int(window_size))
    points: List[Dict[str, Any]] = []

    for start in range(0, len(history), step):
        items = history[start : start + step]
        centroid_values = [item["centroid"] for item in items if item.get("centroid") is not None]
        if not centroid_values:
            continue

        target_row = sum(value[0] for value in centroid_values) / len(centroid_values)
        target_col = sum(value[1] for value in centroid_values) / len(centroid_values)

        representative_item: Optional[Dict[str, Any]] = None
        for item in reversed(items):
            if item.get("polluted_ids"):
                representative_item = item
                break
        if representative_item is None:
            continue

        snapped = snap_centroid_to_polluted_point(
            set(representative_item["polluted_ids"]),
            target_row,
            target_col,
        )
        if snapped is None:
            continue

        point_id, snapped_row, snapped_col = snapped
        start_label = _format_axis_date_label(str(items[0].get("switch_label") or items[0].get("date") or "未识别日期"))
        end_label = _format_axis_date_label(
            str(items[-1].get("switch_label") or items[-1].get("date") or "未识别日期")
        )
        representative_label = _format_axis_date_label(
            str(representative_item.get("switch_label") or representative_item.get("date") or end_label)
        )
        period_label = start_label if start_label == end_label else f"{start_label} ~ {end_label}"

        points.append(
            {
                "index_start": start,
                "index_end": start + len(items) - 1,
                "window_span": len(items),
                "window_start_label": start_label,
                "window_end_label": end_label,
                "period_label": period_label,
                "representative_label": representative_label,
                "point_id": point_id,
                "row": snapped_row,
                "col": snapped_col,
                "target_row": target_row,
                "target_col": target_col,
                "polluted_count": int(representative_item["polluted_count"]),
                "cluster_count": int(representative_item["cluster_count"]),
            }
        )

    return points


def point_distance(point_a: int, point_b: int) -> int:
    row_a, col_a = id_to_row_col(point_a)
    row_b, col_b = id_to_row_col(point_b)
    return abs(row_a - row_b) + abs(col_a - col_b)


def iter_neighbor_ids(point_id: int, radius: int = 1) -> Set[int]:
    row, col = id_to_row_col(point_id)
    neighbors: Set[int] = set()
    for row_offset in range(-radius, radius + 1):
        for col_offset in range(-radius, radius + 1):
            if row_offset == 0 and col_offset == 0:
                continue
            if abs(row_offset) + abs(col_offset) > radius:
                continue
            next_row = row + row_offset
            next_col = col + col_offset
            if 1 <= next_row <= ROWS and 1 <= next_col <= COLS:
                neighbors.add(row_col_to_id(next_row, next_col))
    return neighbors


def count_clusters(point_ids: Set[int]) -> int:
    remaining = set(point_ids)
    clusters = 0
    while remaining:
        clusters += 1
        stack = [remaining.pop()]
        while stack:
            current = stack.pop()
            connected = iter_neighbor_ids(current) & remaining
            if connected:
                stack.extend(connected)
                remaining -= connected
    return clusters


def build_report_history(
    report_options: List[Tuple[str, str]], report_store: Dict[str, Dict[str, Any]]
) -> List[Dict[str, Any]]:
    history: List[Dict[str, Any]] = []
    previous_ids: Set[int] = set()

    for index, (report_key, switch_label) in enumerate(report_options):
        report = report_store.get(report_key, {})
        report_text = str(report.get("report_text", ""))
        parsed = parse_report_text(report_text)
        polluted_ids = get_report_polluted_ids(report)
        obstacle_ids = normalize_point_set(report.get("obstacle_ids"))
        pending_ids = normalize_point_set(report.get("pending_ids"))
        accessible_count = max(TOTAL - len(obstacle_ids | pending_ids), 1)
        centroid = compute_centroid(polluted_ids)
        added_ids = polluted_ids - previous_ids if history else set(polluted_ids)
        removed_ids = previous_ids - polluted_ids if history else set()

        history.append(
            {
                "index": index,
                "report_key": report_key,
                "switch_label": switch_label,
                "date": str(report.get("date") or "未识别日期"),
                "file_name": str(report.get("file_name", "-")),
                "polluted_ids": polluted_ids,
                "obstacle_ids": obstacle_ids,
                "pending_ids": pending_ids,
                "glove_names": parsed.glove_names,
                "polluted_count": len(polluted_ids),
                "glove_count": len(parsed.glove_names),
                "accessible_count": accessible_count,
                "pollution_ratio": len(polluted_ids) / accessible_count,
                "added_ids": added_ids,
                "removed_ids": removed_ids,
                "cluster_count": count_clusters(polluted_ids),
                "centroid": centroid,
            }
        )
        previous_ids = polluted_ids

    return history


def build_point_history_stats(history: List[Dict[str, Any]]) -> Dict[int, Dict[str, Any]]:
    point_ids = set()
    for item in history:
        point_ids |= set(item["polluted_ids"])

    stats: Dict[int, Dict[str, Any]] = {}
    total_reports = max(len(history), 1)

    for point_id in point_ids:
        days_present: List[int] = []
        max_streak = 0
        current_streak = 0
        rolling_streak = 0
        for index, item in enumerate(history):
            if point_id in item["polluted_ids"]:
                days_present.append(index)
                rolling_streak += 1
                max_streak = max(max_streak, rolling_streak)
            else:
                rolling_streak = 0
        for item in reversed(history):
            if point_id in item["polluted_ids"]:
                current_streak += 1
            else:
                break

        stats[point_id] = {
            "occurrences": len(days_present),
            "frequency_ratio": len(days_present) / total_reports,
            "days_present": days_present,
            "first_seen_index": days_present[0],
            "last_seen_index": days_present[-1],
            "max_streak": max_streak,
            "current_streak": current_streak,
        }
    return stats


def forecast_next_pollution_count(
    history: List[Dict[str, Any]], config: Optional[Dict[str, float]] = None
) -> int:
    if not history:
        return 0
    params = merge_analysis_config(config)
    counts = [int(item["polluted_count"]) for item in history]
    if len(counts) == 1:
        return counts[0]

    window_size = max(2, min(int(params["forecast_window"]), len(counts)))
    window = counts[-window_size:]
    weights = list(range(1, len(window) + 1))
    weighted_avg = sum(value * weight for value, weight in zip(window, weights)) / sum(weights)
    diffs = [window[idx] - window[idx - 1] for idx in range(1, len(window))]
    drift = sum(diffs) / len(diffs) if diffs else 0.0
    blend = normalize_weights(
        {
            "smooth": params["forecast_smooth_weight"],
            "trend": params["forecast_trend_weight"],
        }
    )
    trend_projection = window[-1] + drift * params["forecast_drift_scale"]
    forecast = round(weighted_avg * blend["smooth"] + trend_projection * blend["trend"])
    return max(0, min(TOTAL, forecast))


def build_frequency_map(history: List[Dict[str, Any]]) -> Dict[int, int]:
    counts: Counter[int] = Counter()
    for item in history:
        counts.update(item["polluted_ids"])
    return dict(counts)


def rows_to_csv_bytes(rows: List[Dict[str, Any]]) -> bytes:
    if not rows:
        return "\ufeff".encode("utf-8")

    buffer = io.StringIO()
    fieldnames = list(rows[0].keys())
    writer = csv.DictWriter(buffer, fieldnames=fieldnames, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        normalized_row = {key: ("" if row.get(key) is None else str(row.get(key))) for key in fieldnames}
        writer.writerow(normalized_row)
    return ("\ufeff" + buffer.getvalue()).encode("utf-8")


def build_risk_candidates(
    history: List[Dict[str, Any]], config: Optional[Dict[str, float]] = None
) -> List[Dict[str, Any]]:
    if not history:
        return []

    params = merge_analysis_config(config)
    weights = normalize_weights(
        {
            "frequency": params["risk_weight_frequency"],
            "recent": params["risk_weight_recent"],
            "streak": params["risk_weight_streak"],
            "neighbor": params["risk_weight_neighbor"],
            "regional": params["risk_weight_regional"],
        }
    )
    stats = build_point_history_stats(history)
    latest = history[-1]
    latest_polluted_ids = set(latest["polluted_ids"])
    recent_history = history[-min(3, len(history)) :]
    recent_reports = max(len(recent_history), 1)
    excluded_ids = set(latest["obstacle_ids"]) | set(latest["pending_ids"])
    candidates: List[Dict[str, Any]] = []

    for point_id in range(1, TOTAL + 1):
        if point_id in excluded_ids:
            continue

        point_stats = stats.get(
            point_id,
            {
                "occurrences": 0,
                "frequency_ratio": 0.0,
                "max_streak": 0,
                "current_streak": 0,
                "days_present": [],
            },
        )
        recent_occurrences = sum(1 for item in recent_history if point_id in item["polluted_ids"])
        recent_ratio = recent_occurrences / recent_reports
        neighbors = iter_neighbor_ids(point_id)
        near_neighbors = iter_neighbor_ids(point_id, radius=2) - {point_id}
        immediate_support = (
            len(neighbors & latest_polluted_ids) / max(len(neighbors), 1)
            if neighbors
            else 0.0
        )
        regional_support = (
            len(near_neighbors & latest_polluted_ids) / max(len(near_neighbors), 1)
            if near_neighbors
            else 0.0
        )
        persistence_ratio = min(point_stats["current_streak"] / max(len(history), 1), 1.0)

        score = (
            weights["frequency"] * point_stats["frequency_ratio"]
            + weights["recent"] * recent_ratio
            + weights["streak"] * persistence_ratio
            + weights["neighbor"] * immediate_support
            + weights["regional"] * regional_support
        )
        if point_id in latest_polluted_ids:
            score += params["risk_bonus_current"]
        elif immediate_support >= 0.34:
            score += params["risk_bonus_spread"]

        score = max(0.0, min(1.0, score))
        if score < params["risk_threshold"]:
            continue

        row, col = id_to_row_col(point_id)
        candidates.append(
            {
                "point_id": point_id,
                "row": row,
                "col": col,
                "score": score,
                "current_state": "当前污染" if point_id in latest_polluted_ids else "潜在扩散",
                "frequency_ratio": point_stats["frequency_ratio"],
                "recent_ratio": recent_ratio,
                "current_streak": point_stats["current_streak"],
                "neighbor_support": immediate_support,
                "regional_support": regional_support,
            }
        )

    candidates.sort(key=lambda item: (-item["score"], item["point_id"]))
    return candidates


def build_source_candidates(
    history: List[Dict[str, Any]], config: Optional[Dict[str, float]] = None
) -> List[Dict[str, Any]]:
    if not history:
        return []

    params = merge_analysis_config(config)
    weights = normalize_weights(
        {
            "first": params["source_weight_first"],
            "persistence": params["source_weight_persistence"],
            "spread": params["source_weight_spread"],
            "frequency": params["source_weight_frequency"],
            "centroid": params["source_weight_centroid"],
        }
    )
    stats = build_point_history_stats(history)
    frequency_map = build_frequency_map(history)
    total_reports = max(len(history), 1)
    weighted_total = max(sum(frequency_map.values()), 1)

    centroid_row = sum(id_to_row_col(point_id)[0] * count for point_id, count in frequency_map.items()) / weighted_total
    centroid_col = sum(id_to_row_col(point_id)[1] * count for point_id, count in frequency_map.items()) / weighted_total
    first_seen_lookup = {point_id: value["first_seen_index"] for point_id, value in stats.items()}

    candidates: List[Dict[str, Any]] = []
    for point_id, point_stats in stats.items():
        row, col = id_to_row_col(point_id)
        first_seen_index = point_stats["first_seen_index"]
        first_seen_score = 1.0 if total_reports == 1 else 1.0 - first_seen_index / (total_reports - 1)
        persistence_score = point_stats["max_streak"] / total_reports
        frequency_score = point_stats["frequency_ratio"]

        influence_targets = 0
        for other_point_id, other_first_seen_index in first_seen_lookup.items():
            if other_point_id == point_id or other_first_seen_index <= first_seen_index:
                continue
            if point_distance(point_id, other_point_id) <= int(params["source_spread_radius"]):
                influence_targets += 1
        spread_score = min(influence_targets / max(params["source_spread_norm"], 1.0), 1.0)

        centroid_distance = abs(row - centroid_row) + abs(col - centroid_col)
        centroid_score = max(0.0, 1.0 - centroid_distance / max(params["source_centroid_norm"], 1.0))

        score = (
            weights["first"] * first_seen_score
            + weights["persistence"] * persistence_score
            + weights["spread"] * spread_score
            + weights["frequency"] * frequency_score
            + weights["centroid"] * centroid_score
        )
        score = max(0.0, min(1.0, score))

        explanation_parts = [
            f"首次出现第 {first_seen_index + 1} 天",
            f"最长连续 {point_stats['max_streak']} 天",
            f"历史出现 {point_stats['occurrences']} 次",
        ]
        if influence_targets:
            explanation_parts.append(f"后续带动周边 {influence_targets} 个点扩散")

        candidates.append(
            {
                "point_id": point_id,
                "row": row,
                "col": col,
                "score": score,
                "first_seen_index": first_seen_index,
                "max_streak": point_stats["max_streak"],
                "occurrences": point_stats["occurrences"],
                "spread_targets": influence_targets,
                "explanation": "；".join(explanation_parts),
            }
        )

    candidates.sort(key=lambda item: (-item["score"], item["first_seen_index"], item["point_id"]))
    return candidates


def default_analysis_config() -> Dict[str, float]:
    return {
        "forecast_window": 4.0,
        "forecast_smooth_weight": 0.78,
        "forecast_trend_weight": 0.22,
        "forecast_drift_scale": 1.0,
        "risk_weight_frequency": 0.30,
        "risk_weight_recent": 0.24,
        "risk_weight_streak": 0.18,
        "risk_weight_neighbor": 0.18,
        "risk_weight_regional": 0.10,
        "risk_bonus_current": 0.08,
        "risk_bonus_spread": 0.05,
        "risk_threshold": 0.18,
        "source_weight_first": 0.28,
        "source_weight_persistence": 0.24,
        "source_weight_spread": 0.22,
        "source_weight_frequency": 0.16,
        "source_weight_centroid": 0.10,
        "source_spread_radius": 3.0,
        "source_spread_norm": 8.0,
        "source_centroid_norm": 18.0,
    }


def merge_analysis_config(overrides: Optional[Dict[str, float]] = None) -> Dict[str, float]:
    config = default_analysis_config()
    if overrides:
        config.update(overrides)
    return config


def normalize_weights(weights: Dict[str, float]) -> Dict[str, float]:
    total = sum(max(value, 0.0) for value in weights.values())
    if total <= 0:
        uniform = 1.0 / max(len(weights), 1)
        return {key: uniform for key in weights}
    return {key: max(value, 0.0) / total for key, value in weights.items()}


def build_playback_labels(history: List[Dict[str, Any]]) -> List[str]:
    return [f"第 {index + 1} 期 | {item['date']}" for index, item in enumerate(history)]


def build_analytics_snapshot(
    history: List[Dict[str, Any]], end_index: int, config: Optional[Dict[str, float]] = None
) -> Dict[str, Any]:
    clipped_history = history[: end_index + 1]
    latest = clipped_history[-1]
    return {
        "history": clipped_history,
        "latest": latest,
        "forecast_count": forecast_next_pollution_count(clipped_history, config),
        "frequency_map": build_frequency_map(clipped_history),
        "point_stats": build_point_history_stats(clipped_history),
        "risk_candidates": build_risk_candidates(clipped_history, config),
        "source_candidates": build_source_candidates(clipped_history, config),
    }


def render_playback_controls(history: List[Dict[str, Any]], key: str, title: str) -> int:
    import streamlit as st

    labels = build_playback_labels(history)
    max_index = len(labels) - 1
    index_key = f"{key}_index"
    autoplay_key = f"{key}_autoplay"
    finished_key = f"{key}_finished"
    speed_key = f"{key}_speed"
    owner_key = "analytics_autoplay_owner"
    speed_options = {
        "0.5x": 1.35,
        "1x": 0.8,
        "2x": 0.4,
    }

    if index_key not in st.session_state:
        st.session_state[index_key] = max_index
    st.session_state[index_key] = max(0, min(int(st.session_state[index_key]), max_index))
    if autoplay_key not in st.session_state:
        st.session_state[autoplay_key] = False
    if finished_key not in st.session_state:
        st.session_state[finished_key] = False
    if speed_key not in st.session_state:
        st.session_state[speed_key] = "1x"

    status_markup = (
        '<div class="playback-head-status">自动播放中</div>' if st.session_state[autoplay_key] else ""
    )
    st.markdown(
        f"""
        <div class="playback-head">
            <div class="playback-head-title">{html.escape(title)}</div>
            {status_markup}
        </div>
        """,
        unsafe_allow_html=True,
    )
    render_playback_progress(
        st.session_state[index_key],
        len(labels),
        labels[st.session_state[index_key]],
    )

    previous_index = int(st.session_state[index_key])
    selected_label = st.select_slider(
        title,
        options=labels,
        value=labels[st.session_state[index_key]],
        key=f"{key}_slider",
        label_visibility="collapsed",
    )
    st.session_state[index_key] = labels.index(selected_label)
    if not st.session_state[autoplay_key] and st.session_state[index_key] != previous_index:
        st.session_state[finished_key] = False

    control_col, speed_col = st.columns([4.3, 1.2], gap="medium")
    with speed_col:
        st.selectbox(
            "播放速度",
            options=list(speed_options.keys()),
            key=speed_key,
            label_visibility="collapsed",
        )

    b1, b2, b3, b4, b5 = control_col.columns(5)
    with b1:
        if st.button("上一步", key=f"{key}_prev", disabled=st.session_state[index_key] <= 0):
            st.session_state[index_key] -= 1
            st.session_state[autoplay_key] = False
            st.session_state[finished_key] = False
            if st.session_state.get(owner_key) == key:
                st.session_state[owner_key] = ""
            st.rerun()
    with b2:
        if st.button("下一步", key=f"{key}_next", disabled=st.session_state[index_key] >= max_index):
            st.session_state[index_key] += 1
            st.session_state[autoplay_key] = False
            st.session_state[finished_key] = False
            if st.session_state.get(owner_key) == key:
                st.session_state[owner_key] = ""
            st.rerun()
    with b3:
        if st.button("自动播放", key=f"{key}_play", disabled=max_index <= 0):
            st.session_state[autoplay_key] = True
            st.session_state[finished_key] = False
            st.session_state[owner_key] = key
            if st.session_state[index_key] >= max_index:
                st.session_state[index_key] = 0
            st.rerun()
    with b4:
        if st.button("停止", key=f"{key}_stop", disabled=not st.session_state[autoplay_key]):
            st.session_state[autoplay_key] = False
            st.session_state[finished_key] = False
            if st.session_state.get(owner_key) == key:
                st.session_state[owner_key] = ""
            st.rerun()
    with b5:
        if st.button("回到最新", key=f"{key}_latest", disabled=st.session_state[index_key] >= max_index):
            st.session_state[index_key] = max_index
            st.session_state[autoplay_key] = False
            st.session_state[finished_key] = False
            if st.session_state.get(owner_key) == key:
                st.session_state[owner_key] = ""
            st.rerun()

    if (
        st.session_state.get(autoplay_key)
        and st.session_state.get(owner_key) == key
        and st.session_state[index_key] < max_index
    ):
        time.sleep(speed_options.get(str(st.session_state.get(speed_key)), 0.8))
        st.session_state[index_key] += 1
        st.rerun()

    if st.session_state.get(autoplay_key) and st.session_state[index_key] >= max_index:
        st.session_state[autoplay_key] = False
        st.session_state[finished_key] = True
        if st.session_state.get(owner_key) == key:
            st.session_state[owner_key] = ""

    return int(st.session_state[index_key])


def build_source_diffusion_figure(snapshot: Dict[str, Any], config: Optional[Dict[str, float]] = None):
    import plotly.graph_objects as go

    params = merge_analysis_config(config)
    source_candidates = list(snapshot.get("source_candidates", []))
    point_stats = snapshot.get("point_stats", {})
    if not source_candidates or not point_stats:
        fig = go.Figure()
        fig.update_layout(
            title="扩散链路回放",
            height=520,
            margin={"l": 18, "r": 14, "t": 56, "b": 24},
            paper_bgcolor="rgba(255,255,255,0)",
            plot_bgcolor="rgba(255,255,255,0.92)",
            annotations=[
                {
                    "text": "当前回放期内没有足够的扩散链路",
                    "xref": "paper",
                    "yref": "paper",
                    "x": 0.5,
                    "y": 0.5,
                    "showarrow": False,
                    "font": {"size": 16, "color": "#47698b"},
                }
            ],
        )
        return fig

    radius = int(params["source_spread_radius"])
    palette = ["#e14b43", "#2e6ea9", "#3ea66b"]
    ranked_sources = source_candidates[:3]
    edges: List[Tuple[int, int, str]] = []
    node_ids: Set[int] = set()

    for index, source in enumerate(ranked_sources):
        source_id = int(source["point_id"])
        source_first = int(source["first_seen_index"])
        color = palette[index % len(palette)]
        targets: List[Tuple[int, int, int]] = []
        for target_id, target_stats in point_stats.items():
            if target_id == source_id:
                continue
            target_first = int(target_stats["first_seen_index"])
            if target_first <= source_first:
                continue
            distance = point_distance(source_id, target_id)
            if distance <= radius + 1:
                targets.append((target_id, target_first, distance))
        targets.sort(key=lambda item: (item[1], item[2], item[0]))
        for target_id, _, _ in targets[:10]:
            edges.append((source_id, target_id, color))
            node_ids.add(source_id)
            node_ids.add(target_id)

    if not edges:
        node_ids = {int(item["point_id"]) for item in ranked_sources}

    ordered_nodes = sorted(
        node_ids,
        key=lambda point_id: (
            int(point_stats.get(point_id, {}).get("first_seen_index", 999)),
            id_to_row_col(point_id)[0],
            id_to_row_col(point_id)[1],
            point_id,
        ),
    )
    y_positions = {point_id: len(ordered_nodes) - idx for idx, point_id in enumerate(ordered_nodes)}
    node_source_rank = {int(item["point_id"]): rank + 1 for rank, item in enumerate(ranked_sources)}

    fig = go.Figure()
    for source_id, target_id, color in edges:
        source_stats = point_stats[source_id]
        target_stats = point_stats[target_id]
        fig.add_trace(
            go.Scatter(
                x=[source_stats["first_seen_index"] + 1, target_stats["first_seen_index"] + 1],
                y=[y_positions[source_id], y_positions[target_id]],
                mode="lines",
                line={"color": color, "width": 2.4},
                hovertemplate=(
                    f"扩散链路<br>ID {source_id} -> ID {target_id}<br>"
                    f"源点首次出现: 第 {source_stats['first_seen_index'] + 1} 期<br>"
                    f"目标首次出现: 第 {target_stats['first_seen_index'] + 1} 期<extra></extra>"
                ),
                showlegend=False,
            )
        )

    marker_xs: List[int] = []
    marker_ys: List[int] = []
    marker_sizes: List[float] = []
    marker_colors: List[str] = []
    marker_symbols: List[str] = []
    marker_texts: List[str] = []
    marker_hover: List[str] = []

    for point_id in ordered_nodes:
        stats = point_stats[point_id]
        row, col = id_to_row_col(point_id)
        marker_xs.append(stats["first_seen_index"] + 1)
        marker_ys.append(y_positions[point_id])
        marker_sizes.append(14 + stats["occurrences"] * 3)
        if point_id in node_source_rank:
            marker_colors.append("#102b47")
            marker_symbols.append("diamond")
            marker_texts.append(f"S{node_source_rank[point_id]}")
        else:
            marker_colors.append("#6fa0cf")
            marker_symbols.append("circle")
            marker_texts.append(str(point_id))
        marker_hover.append(
            f"ID {point_id}<br>行{row} 列{col}<br>首次出现: 第 {stats['first_seen_index'] + 1} 期"
            f"<br>出现次数: {stats['occurrences']}<br>最长连续: {stats['max_streak']}"
        )

    fig.add_trace(
        go.Scatter(
            x=marker_xs,
            y=marker_ys,
            mode="markers+text",
            text=marker_texts,
            textposition="middle center",
            marker={
                "size": marker_sizes,
                "color": marker_colors,
                "symbol": marker_symbols,
                "line": {"width": 2, "color": "#ffffff"},
            },
            textfont={"size": 10, "color": "#ffffff"},
            hovertext=marker_hover,
            hovertemplate="%{hovertext}<extra></extra>",
            showlegend=False,
        )
    )

    fig.update_layout(
        title="疑似源点扩散链路回放",
        height=560,
        margin={"l": 18, "r": 14, "t": 56, "b": 24},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.92)",
    )
    fig.update_xaxes(
        title_text="首次出现期次",
        tickmode="array",
        tickvals=list(range(1, len(snapshot["history"]) + 1)),
        gridcolor="#e4edf7",
        rangemode="tozero",
    )
    fig.update_yaxes(
        title_text="扩散链路节点",
        tickmode="array",
        tickvals=[y_positions[point_id] for point_id in ordered_nodes],
        ticktext=[f"ID {point_id}" for point_id in ordered_nodes],
        gridcolor="#eef4fb",
    )
    return fig


def render_report_switcher(
    title: str, switch_id: str, options: List[Tuple[str, str]], active_key: str
) -> None:
    import streamlit as st

    if not options:
        return

    if title:
        st.caption(title)
    chunk_size = 10
    for start in range(0, len(options), chunk_size):
        row_items = options[start : start + chunk_size]
        cols = st.columns(len(row_items), gap="small")
        for idx, (option_key, label) in enumerate(row_items):
            button_type = "primary" if option_key == active_key else "secondary"
            if cols[idx].button(
                label,
                key=f"{switch_id}_{start + idx}",
                use_container_width=True,
                type=button_type,
            ):
                st.session_state.active_report_key = option_key
                st.rerun()


def default_obstacle_ids() -> Set[int]:
    ids: Set[int] = set()
    for col in range(1, COLS + 1):
        for row in range(1, ROWS + 1):
            obstacle = (
                col in {1, 28}
                or row in {1, 30}
                or (7 <= row <= 10 and 2 <= col <= 22)
                or (21 <= row <= 24 and 2 <= col <= 22)
                or (29 <= row <= 30 and 3 <= col <= 23)
            )
            if obstacle:
                ids.add(row_col_to_id(row, col))
    return ids


def obstacle_ids_from_mask(mask_bytes: bytes) -> Set[int]:
    """Parse obstacle IDs from a 30x28 mask image (dark pixels are obstacles)."""
    try:
        from PIL import Image  # type: ignore
    except Exception as exc:
        raise ValueError("读取掩膜图需要 pillow，请先安装依赖。") from exc

    img = Image.open(io.BytesIO(mask_bytes)).convert("L").resize((COLS, ROWS))
    px = img.load()

    obstacle: Set[int] = set()
    for row in range(1, ROWS + 1):
        for col in range(1, COLS + 1):
            if px[col - 1, row - 1] < 128:
                obstacle.add(row_col_to_id(row, col))
    return obstacle


def apply_selection_action(
    action: str,
    selected_ids: Set[int],
    obstacle_ids: Set[int],
    pending_ids: Set[int],
) -> Tuple[Set[int], Set[int]]:
    new_obstacle = set(obstacle_ids)
    new_pending = set(pending_ids)

    if action == "set_obstacle":
        new_obstacle |= selected_ids
        new_pending -= selected_ids
    elif action == "set_pending":
        new_pending |= selected_ids
        new_obstacle -= selected_ids
    elif action == "clear":
        new_obstacle -= selected_ids
        new_pending -= selected_ids

    return new_obstacle, new_pending


def _event_to_dict(selection_event: Any) -> Dict[str, Any]:
    if isinstance(selection_event, dict):
        return selection_event
    if hasattr(selection_event, "to_dict"):
        converted = selection_event.to_dict()
        if isinstance(converted, dict):
            return converted
    selection = getattr(selection_event, "selection", None)
    if isinstance(selection, dict):
        return {"selection": selection}
    if hasattr(selection, "to_dict"):
        converted = selection.to_dict()
        if isinstance(converted, dict):
            return {"selection": converted}
    return {}


def has_selection_payload(selection_event: Any) -> bool:
    event = _event_to_dict(selection_event)
    if not event:
        return False
    selection = event.get("selection")
    if isinstance(selection, dict) and "points" in selection:
        return True
    return isinstance(event.get("points"), list)


def extract_selected_ids(selection_event: Any) -> Set[int]:
    event = _event_to_dict(selection_event)
    if not event:
        return set()

    points = []
    if isinstance(event.get("selection"), dict):
        points = event.get("selection", {}).get("points", [])
    if not points and isinstance(event.get("points"), list):
        points = event.get("points", [])

    ids: Set[int] = set()
    for point in points:
        custom = point.get("customdata")
        if isinstance(custom, list) and custom:
            custom = custom[0]
        if isinstance(custom, int) and 1 <= custom <= TOTAL:
            ids.add(custom)
    return ids


def build_plotly_grid(
    polluted_ids: Set[int],
    obstacle_ids: Set[int],
    pending_ids: Set[int],
    selected_ids: Optional[Set[int]] = None,
):
    import plotly.graph_objects as go

    polluted_fill = "#ff4d4f"
    polluted_text = "#ffffff"
    selected_fill = "#2e6ea9"
    selected_text = "#ffffff"
    staged_ids = set(selected_ids or set())

    xs: List[int] = []
    ys: List[int] = []
    ids: List[int] = []
    colors: List[str] = []
    text_colors: List[str] = []
    for col in range(1, COLS + 1):
        for row in range(1, ROWS + 1):
            point_id = row_col_to_id(row, col)
            xs.append(col)
            ys.append(ROWS - row + 1)
            ids.append(point_id)
            if point_id in staged_ids:
                colors.append(selected_fill)
                text_colors.append(selected_text)
            elif point_id in obstacle_ids:
                colors.append("#8d9096")
                text_colors.append("#ffffff")
            elif point_id in pending_ids:
                colors.append("#ffffff")
                text_colors.append("#333333")
            elif point_id in polluted_ids:
                colors.append(polluted_fill)
                text_colors.append(polluted_text)
            else:
                colors.append("#d2f4d2")
                text_colors.append("#1b5e20")

    fig = go.Figure()
    fig.add_trace(
        go.Scatter(
            x=xs,
            y=ys,
            mode="markers+text",
            text=[str(point_id) for point_id in ids],
            textposition="middle center",
            textfont={"size": 9, "color": text_colors},
            customdata=ids,
            marker={
                "symbol": "square",
                "size": 21,
                "color": colors,
                "line": {"width": 0.5, "color": "#dddddd"},
            },
            selected={"marker": {"opacity": 1.0, "size": 22}},
            unselected={"marker": {"opacity": 0.95}},
            hovertemplate="ID=%{customdata}<extra></extra>",
        )
    )

    fig.update_layout(
        dragmode="select",
        uirevision="grid-static",
        newselection={"line": {"color": "#2e6ea9", "width": 2}},
        activeselection={"fillcolor": "rgba(46, 110, 169, 0.18)", "opacity": 0.35},
        transition={"duration": 120, "easing": "cubic-in-out"},
        title="可多次拖拽累积选中点位，最后点击下方按钮确认",
        xaxis={"range": [0.5, COLS + 0.5], "showgrid": False, "visible": False},
        yaxis={
            "range": [0.5, ROWS + 0.5],
            "showgrid": False,
            "visible": False,
            "scaleanchor": "x",
            "scaleratio": 1,
        },
        margin={"l": 16, "r": 16, "t": 40, "b": 12},
        height=840,
    )
    return fig


def build_result_snapshot_grid(
    polluted_ids: Set[int],
    obstacle_ids: Set[int],
    pending_ids: Set[int],
):
    fig = build_plotly_grid(polluted_ids, obstacle_ids, pending_ids, selected_ids=None)
    fig.update_layout(
        title="当前日期地图结果",
        dragmode="pan",
        margin={"l": 16, "r": 16, "t": 44, "b": 12},
        height=840,
    )
    return fig


def build_analytics_date_map_figure(item: Dict[str, Any]):
    fig = build_plotly_grid(
        set(item["polluted_ids"]),
        set(item["obstacle_ids"]),
        set(item["pending_ids"]),
        selected_ids=None,
    )
    date_label = _format_axis_date_label(str(item.get("switch_label") or item.get("date") or "未识别日期"))
    fig.update_layout(
        title=f"日期地图快照 | {date_label}",
        dragmode=False,
        margin={"l": 12, "r": 12, "t": 52, "b": 12},
        height=560,
    )
    return fig


def inject_styles() -> None:
    import streamlit as st

    st.markdown(
        """
        <style>
        :root {
            --ink-900: #102b47;
            --ink-700: #214a71;
            --line-200: #bdd1e5;
            --line-300: #a9c2dd;
            --surface-0: #ffffff;
            --surface-1: #f4f8fd;
            --surface-2: #e9f1fa;
            --brand-500: #2e6ea9;
            --brand-600: #245c90;
            --brand-700: #1f4e7a;
            --section-title-size: 1.42rem;
            --result-title-size: 1.65rem;
            --compact-panel-h: 98px;
            --map-panel-h: 840px;
            --result-section-offset: 16px;
            /* 这里控制右侧三块明细区彼此之间的上下距离。
               如果你想把“当前污染点位 ID”下面的明细框和下面一块再拉开一点，
               就把这个值调大。 */
            --detail-gap: 20px;
            /* 这里控制“当前无污染点位”指标卡和下方“当前污染点位 ID”标题之间的距离。
               想调你截图箭头指的那一段空白，就改这个值。 */
            --metrics-to-detail-gap: 18px;
            --result-top-offset: 40px;
            /* 布局微调入口：
               1. `--result-top-offset` 控制“识别结果”相对地图上边沿的对齐
               2. `--detail-total-h` 控制右侧三块明细区总高度
               3. `--detail-gap` 控制右侧明细区卡片之间的垂直间距
               4. `--compact-panel-h` 控制左侧“当前框选”白框高度
               5. `--map-panel-h` 控制地图整体高度 */
            --detail-total-h: calc(var(--map-panel-h) - 56px);
        }
        [data-testid="stAppViewContainer"] {
            background:
                radial-gradient(1200px 420px at 8% -10%, #ffffffaa 0%, #ffffff00 55%),
                radial-gradient(1000px 380px at 100% 0%, #c9dcf3aa 0%, #c9dcf300 52%),
                linear-gradient(180deg, #e9f2fc 0%, #d4e3f4 100%);
        }
        .main .block-container {
            max-width: 1500px;
            padding-top: 1.25rem;
            padding-bottom: 1.4rem;
        }
        html, body, [class*="css"] {
            font-family: "Segoe UI", "PingFang SC", "Microsoft YaHei", sans-serif;
            color: var(--ink-900);
        }
        h1, h2, h3 {
            color: var(--ink-900);
            letter-spacing: 0.2px;
        }
        h1 {
            font-weight: 760;
            margin-bottom: 0.4rem;
        }
        h1::after {
            content: "";
            display: block;
            width: 84px;
            height: 4px;
            margin-top: 10px;
            border-radius: 999px;
            background: linear-gradient(90deg, #2f6ea8 0%, #6fa0cf 100%);
        }
        h3 {
            font-weight: 680;
            font-size: var(--section-title-size);
        }
        .result-title {
            color: #1f4468;
            font-size: var(--result-title-size);
            font-weight: 700;
            line-height: 1.25;
            letter-spacing: 0.2px;
            margin: 0 0 16px 0;
        }
        .result-section-offset {
            height: var(--result-section-offset);
        }
        .result-top-spacer {
            /* 这里是右侧结果栏的顶部定位器。
               想让“识别结果”和地图上边沿重新对齐，就改这个高度。 */
            height: var(--result-top-offset);
        }
        [data-testid="stPlotlyChart"] {
            background: linear-gradient(180deg, #ffffff 0%, #f6faff 100%);
            border: 1px solid var(--line-300);
            border-radius: 16px;
            padding: 6px 6px 2px 6px;
            box-sizing: border-box;
            overflow: hidden;
            box-shadow:
                0 10px 24px rgba(10, 41, 74, 0.12),
                inset 0 1px 0 rgba(255, 255, 255, 0.8);
        }
        [data-testid="stPlotlyChart"] > div,
        [data-testid="stPlotlyChart"] .js-plotly-plot,
        [data-testid="stPlotlyChart"] .plot-container,
        [data-testid="stPlotlyChart"] .svg-container {
            max-width: 100% !important;
            width: 100% !important;
            box-sizing: border-box;
        }
        [data-testid="stFileUploader"], [data-testid="stTextArea"], [data-testid="stRadio"], [data-testid="stSelectbox"], [data-testid="stSelectSlider"] {
            background:
                radial-gradient(220px 80px at 100% 0%, rgba(255,255,255,0.34) 0%, rgba(255,255,255,0) 70%),
                linear-gradient(180deg, rgba(255,255,255,0.92) 0%, rgba(242,247,253,0.98) 100%);
            border-radius: 16px;
            border: 1px solid rgba(189, 209, 229, 0.9);
            padding: 0.35rem 0.45rem;
            box-shadow:
                0 8px 18px rgba(13, 45, 79, 0.06),
                inset 0 1px 0 rgba(255,255,255,0.82);
        }
        [data-testid="stFileUploaderDropzone"] {
            border-radius: 14px !important;
            border: 1px dashed rgba(47, 110, 169, 0.35) !important;
            background: linear-gradient(180deg, rgba(252,254,255,0.96) 0%, rgba(238,245,252,0.98) 100%) !important;
            padding-top: 1rem !important;
            padding-bottom: 1rem !important;
        }
        [data-testid="stFileUploaderDropzone"]:hover {
            border-color: rgba(38, 96, 151, 0.48) !important;
            background: linear-gradient(180deg, #ffffff 0%, #eef5fd 100%) !important;
        }
        [data-testid="stFileUploaderDropzoneInstructions"] div {
            color: #547595 !important;
        }
        [data-testid="stFileUploader"] small {
            color: #6a87a3 !important;
        }
        [data-testid="stTextArea"] label,
        [data-testid="stRadio"] label[data-baseweb="radio"] > div:last-child,
        [data-testid="stFileUploader"] label {
            color: #294c70 !important;
        }
        [data-baseweb="textarea"] textarea {
            background: linear-gradient(180deg, #ffffff 0%, #f7fbff 100%);
            border-radius: 12px !important;
            padding: 0.75rem 0.9rem !important;
            color: #173c61 !important;
            line-height: 1.6 !important;
        }
        [data-baseweb="textarea"] textarea:focus {
            box-shadow: 0 0 0 2px rgba(46, 110, 169, 0.12) !important;
        }
        [data-baseweb="select"] > div {
            border-radius: 12px !important;
            background: linear-gradient(180deg, #ffffff 0%, #f7fbff 100%) !important;
            border-color: rgba(189, 209, 229, 0.95) !important;
            color: #173c61 !important;
            min-height: 42px !important;
        }
        [data-baseweb="select"] > div:hover {
            border-color: rgba(92, 138, 182, 0.6) !important;
        }
        [data-testid="stSelectSlider"] label {
            color: #294c70 !important;
        }
        [data-testid="stRadio"] [role="radiogroup"] {
            gap: 10px;
        }
        [data-testid="stRadio"] label[data-baseweb="radio"] {
            min-height: 42px;
            border-radius: 999px;
            border: 1px solid rgba(47, 110, 169, 0.18);
            background:
                radial-gradient(180px 54px at 100% 0%, rgba(112, 167, 217, 0.18) 0%, rgba(112, 167, 217, 0) 65%),
                linear-gradient(180deg, rgba(255,255,255,0.94) 0%, rgba(234,242,251,0.98) 100%);
            padding: 0 14px;
            box-shadow:
                0 5px 12px rgba(17, 54, 91, 0.06),
                inset 0 1px 0 rgba(255,255,255,0.78);
            transition: all 0.16s ease;
            display: flex;
            align-items: center;
        }
        [data-testid="stRadio"] label[data-baseweb="radio"] > div:last-child {
            display: flex;
            align-items: center;
            line-height: 1.34;
            min-height: 100%;
            padding-bottom: 1px;
        }
        [data-testid="stRadio"] label[data-baseweb="radio"]:has(input:checked) {
            border-color: rgba(36, 92, 144, 0.34);
            background: linear-gradient(180deg, #3b79b5 0%, #2f6ea9 55%, #245b90 100%);
            box-shadow:
                0 8px 16px rgba(19, 60, 101, 0.18),
                inset 0 1px 0 rgba(255,255,255,0.22);
        }
        [data-testid="stRadio"] label[data-baseweb="radio"]:has(input:checked) > div:last-child {
            color: #ffffff !important;
            font-weight: 720;
        }
        [data-testid="stExpander"] {
            border-radius: 16px;
            border: 1px solid rgba(39, 92, 145, 0.14);
            background:
                radial-gradient(260px 100px at 100% 0%, rgba(255,255,255,0.32) 0%, rgba(255,255,255,0) 70%),
                linear-gradient(180deg, rgba(255,255,255,0.88) 0%, rgba(241,247,253,0.98) 100%);
            box-shadow:
                0 10px 22px rgba(16, 45, 79, 0.06),
                inset 0 1px 0 rgba(255,255,255,0.74);
            overflow: hidden;
        }
        [data-testid="stExpander"] summary {
            background: transparent !important;
            color: #20486f !important;
            font-weight: 760 !important;
            padding-top: 0.9rem !important;
            padding-bottom: 0.9rem !important;
        }
        [data-testid="stExpanderDetails"] {
            padding-top: 0.2rem !important;
        }
        .stButton > button {
            min-height: 56px;
            padding: 0.36rem 1.28rem;
            border-radius: 16px;
            font-weight: 760;
            font-size: 1rem;
            letter-spacing: 0.01rem;
            transition: all 0.18s ease;
            cursor: pointer;
            border: 1px solid rgba(114, 160, 205, 0.62);
            background:
                radial-gradient(220px 88px at 100% 0%, rgba(172, 208, 241, 0.28) 0%, rgba(172, 208, 241, 0) 68%),
                linear-gradient(180deg, rgba(255,255,255,0.98) 0%, rgba(240,247,255,0.98) 100%);
            color: #2a5d90;
            box-shadow:
                0 10px 20px rgba(19, 57, 95, 0.10),
                inset 0 1px 0 rgba(255, 255, 255, 0.94);
        }
        .stButton > button[kind="primary"],
        [data-testid="stBaseButton-primary"] {
            border: 1px solid rgba(114, 160, 205, 0.68);
            background:
                radial-gradient(220px 88px at 100% 0%, rgba(172, 208, 241, 0.30) 0%, rgba(172, 208, 241, 0) 68%),
                linear-gradient(180deg, rgba(255,255,255,0.99) 0%, rgba(240,247,255,0.99) 100%);
            color: #245a90;
            box-shadow:
                0 10px 20px rgba(19, 57, 95, 0.12),
                inset 0 1px 0 rgba(255, 255, 255, 0.96);
        }
        .stButton > button[kind="primary"]:hover,
        [data-testid="stBaseButton-primary"]:hover {
            transform: translateY(-1px);
            border-color: rgba(90, 140, 190, 0.88);
            background:
                radial-gradient(220px 88px at 100% 0%, rgba(153, 198, 238, 0.34) 0%, rgba(153, 198, 238, 0) 68%),
                linear-gradient(180deg, #ffffff 0%, #edf6ff 100%);
            color: #1f507f;
            box-shadow:
                0 14px 24px rgba(19, 57, 95, 0.14),
                inset 0 1px 0 rgba(255, 255, 255, 0.98);
        }
        .stButton > button[kind="secondary"],
        [data-testid="stBaseButton-secondary"] {
            border: 1px solid rgba(114, 160, 205, 0.62);
            background:
                radial-gradient(220px 88px at 100% 0%, rgba(172, 208, 241, 0.28) 0%, rgba(172, 208, 241, 0) 68%),
                linear-gradient(180deg, rgba(255,255,255,0.98) 0%, rgba(240,247,255,0.98) 100%);
            color: #2a5d90;
            box-shadow:
                0 10px 20px rgba(19, 57, 95, 0.10),
                inset 0 1px 0 rgba(255, 255, 255, 0.94);
        }
        .stButton > button[kind="secondary"]:hover,
        [data-testid="stBaseButton-secondary"]:hover {
            transform: translateY(-1px);
            border-color: rgba(90, 140, 190, 0.88);
            background:
                radial-gradient(220px 88px at 100% 0%, rgba(153, 198, 238, 0.34) 0%, rgba(153, 198, 238, 0) 68%),
                linear-gradient(180deg, #ffffff 0%, #edf6ff 100%);
            color: #1f507f;
            box-shadow:
                0 14px 24px rgba(19, 57, 95, 0.14),
                inset 0 1px 0 rgba(255, 255, 255, 0.98);
        }
        .stButton > button:active,
        [data-testid^="stBaseButton-"]:active {
            transform: translateY(1px);
            box-shadow:
                0 6px 12px rgba(19, 57, 95, 0.10),
                inset 0 2px 6px rgba(154, 188, 221, 0.26);
        }
        .stButton > button:disabled,
        [data-testid^="stBaseButton-"]:disabled {
            border-color: #c6d8ea;
            background: linear-gradient(180deg, #f8fbfe 0%, #edf3f9 100%);
            color: #9ab0c5;
            box-shadow:
                inset 0 1px 0 rgba(255,255,255,0.88);
            transform: none;
        }
        [data-testid="stDownloadButton"] > button {
            min-height: 56px;
            padding: 0.36rem 1.28rem;
            border-radius: 16px;
            border: 1px solid rgba(114, 160, 205, 0.62);
            background:
                radial-gradient(220px 88px at 100% 0%, rgba(172, 208, 241, 0.28) 0%, rgba(172, 208, 241, 0) 68%),
                linear-gradient(180deg, rgba(255,255,255,0.98) 0%, rgba(240,247,255,0.98) 100%);
            color: #2a5d90;
            font-weight: 760;
            font-size: 1rem;
            box-shadow:
                0 10px 20px rgba(19, 57, 95, 0.10),
                inset 0 1px 0 rgba(255,255,255,0.94);
            cursor: pointer;
        }
        [data-testid="stDownloadButton"] > button:hover {
            transform: translateY(-1px);
            border-color: rgba(90, 140, 190, 0.88);
            background:
                radial-gradient(220px 88px at 100% 0%, rgba(153, 198, 238, 0.34) 0%, rgba(153, 198, 238, 0) 68%),
                linear-gradient(180deg, #ffffff 0%, #edf6ff 100%);
            color: #1f507f;
            box-shadow:
                0 14px 24px rgba(19, 57, 95, 0.14),
                inset 0 1px 0 rgba(255,255,255,0.98);
        }
        [data-testid="stDownloadButton"] {
            margin-bottom: 4px;
        }
        [data-baseweb="tab-list"] {
            gap: 0;
            margin-top: 1.7rem;
            margin-bottom: 1.55rem;
            padding-left: 0;
            border-bottom: none;
            width: 100%;
        }
        [data-baseweb="tab"] {
            height: 70px;
            min-width: 0;
            flex: 1 1 0;
            border-radius: 10px 10px 0 0;
            padding: 0 40px;
            background: linear-gradient(180deg, #fbfdff 0%, #edf4fb 100%);
            border: 1px solid rgba(67, 111, 154, 0.34);
            border-bottom: 1px solid rgba(67, 111, 154, 0.20);
            color: #2b537c;
            box-shadow:
                0 10px 20px rgba(16, 45, 79, 0.10),
                inset 0 1px 0 rgba(255,255,255,0.92);
            cursor: pointer;
            margin-bottom: 0;
            transition: transform 0.15s ease, background 0.15s ease, color 0.15s ease, box-shadow 0.15s ease, border-color 0.15s ease;
        }
        [data-baseweb="tab"] > div,
        [data-baseweb="tab"] span,
        [data-baseweb="tab"] p {
            color: inherit !important;
            font-family: "SimHei", "STHeiti", "Microsoft YaHei", sans-serif !important;
            font-weight: 800 !important;
            font-size: 1.34rem !important;
            letter-spacing: 0.03rem !important;
            line-height: 1.05 !important;
        }
        [data-baseweb="tab"][aria-selected="true"] {
            background: linear-gradient(180deg, #79a9d6 0%, #5d90c4 58%, #4a7db1 100%);
            color: #ffffff;
            border-color: rgba(60, 107, 155, 0.92);
            box-shadow:
                inset 0 4px 0 #d94a43,
                0 14px 24px rgba(18, 53, 89, 0.14);
        }
        [data-baseweb="tab"]:hover {
            background: linear-gradient(180deg, #ffffff 0%, #e3eef9 100%);
            color: #173f65;
            border-color: rgba(45, 94, 146, 0.52);
            transform: translateY(-2px);
            box-shadow:
                0 14px 22px rgba(16, 45, 79, 0.14),
                inset 0 1px 0 rgba(255,255,255,0.92);
        }
        [data-testid="stRadio"] label[data-baseweb="radio"] {
            cursor: pointer;
        }
        .selected-box {
            height: var(--compact-panel-h);
            min-height: var(--compact-panel-h);
            max-height: var(--compact-panel-h);
            overflow-y: auto;
            border: 1px solid var(--line-300);
            border-radius: 13px;
            background: linear-gradient(180deg, #fafdff 0%, #f0f6fe 100%);
            padding: 8px;
            box-shadow:
                inset 0 0 0 1px rgba(255, 255, 255, 0.75),
                0 5px 12px rgba(15, 48, 82, 0.08);
        }
        .selected-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(58px, 1fr));
            gap: 6px;
        }
        .toolbar-selected-box {
            margin-top: 8px;
            margin-bottom: 42px;
        }
        .selected-box.toolbar-selected-box,
        .selected-empty.toolbar-selected-box {
            height: 172px !important;
            min-height: 172px !important;
            max-height: 172px !important;
            box-sizing: border-box;
        }
        .toolbar-selected-grid {
            grid-template-columns: repeat(auto-fill, minmax(66px, 1fr));
            gap: 8px;
        }
        .selected-chip {
            display: block;
            border: 1px solid rgba(156, 186, 216, 0.9);
            background: linear-gradient(180deg, #f8fbfe 0%, #eef4fb 100%);
            color: #23486f;
            border-radius: 8px;
            text-align: center;
            font-size: 0.8rem;
            font-weight: 680;
            line-height: 1.5;
            padding: 3px 8px;
            cursor: default;
        }
        .selected-empty {
            height: var(--compact-panel-h);
            min-height: var(--compact-panel-h);
            max-height: var(--compact-panel-h);
            border: 1px dashed #9fbad5;
            border-radius: 12px;
            background: linear-gradient(180deg, #fbfdff 0%, #f2f7fe 100%);
            color: #3d6286;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 0.88rem;
        }
        .detail-stack {
            /* 这里控制右侧三块明细区的总高度。
               如果最下面的“点位对应行列”没有和左侧“当前框选”白框底边对齐，
               优先微调 `--detail-total-h`。 */
            height: var(--detail-total-h);
            display: grid;
            grid-template-rows: repeat(4, minmax(0, 1fr));
            gap: var(--detail-gap);
        }
        .detail-row {
            display: flex;
            flex-direction: column;
            min-height: 0;
            height: 100%;
        }
        .detail-item {
            border: 1px solid var(--line-300);
            border-radius: 13px;
            background: linear-gradient(180deg, #fafdff 0%, #f0f6fe 100%);
            padding: 6px 8px;
            display: flex;
            flex-direction: column;
            min-height: 0;
            flex: 1;
            height: 100%;
            box-shadow:
                inset 0 0 0 1px rgba(255, 255, 255, 0.75),
                0 5px 12px rgba(15, 48, 82, 0.08);
        }
        .detail-title {
            color: #1f4468;
            font-size: var(--result-title-size);
            font-weight: 700;
            line-height: 1.25;
            letter-spacing: 0.2px;
            /* 这里控制标题文字与它正下方白框之间的距离。
               如果你想把“当前污染点位 ID”这行字和下面白框分开一点，
               就把这个值调大。 */
            margin-bottom: 10px;
            padding-left: 2px;
            display: flex;
            align-items: center;
            gap: 8px;
        }
        .detail-anchor {
            color: #6e8ca9;
            font-size: 0.95rem;
            line-height: 1;
        }
        .detail-body {
            color: #153a5f;
            font-size: 0.8rem;
            line-height: 1.45;
            word-break: break-word;
            overflow-y: auto;
            min-height: 0;
            flex: 1;
        }
        .detail-info-rows {
            display: flex;
            flex-direction: column;
            gap: 0;
        }
        .detail-info-row {
            display: grid;
            grid-template-columns: 92px minmax(0, 1fr);
            gap: 10px;
            align-items: start;
            padding: 7px 2px;
        }
        .detail-info-row + .detail-info-row {
            border-top: 1px solid rgba(189, 209, 229, 0.42);
        }
        .detail-info-label {
            color: #5a7a99;
            font-size: 0.82rem;
            font-weight: 700;
            line-height: 1.5;
        }
        .detail-info-value {
            color: #1c446a;
            font-size: 0.88rem;
            font-weight: 600;
            line-height: 1.56;
            word-break: break-word;
        }
        .metrics-to-detail-gap {
            /* 这是右侧“当前无污染点位”卡片和下方明细标题之间的专用 spacer。
               如果想调整你刚刚指出的那块距离，改 `--metrics-to-detail-gap` 即可。 */
            height: var(--metrics-to-detail-gap);
        }
        .metric-card {
            position: relative;
            overflow: hidden;
            background: linear-gradient(180deg, rgba(255,255,255,0.98) 0%, rgba(244,248,253,0.96) 100%);
            border: 1px solid var(--line-200);
            border-radius: 12px;
            padding: 0.46rem 0.68rem 0.52rem 0.68rem;
            box-shadow:
                0 4px 10px rgba(14, 41, 72, 0.04),
                inset 0 1px 0 rgba(255,255,255,0.86);
            min-height: 90px;
            cursor: default;
        }
        .metric-card::before {
            content: "";
            position: absolute;
            left: 0;
            top: 0;
            width: 100%;
            height: 2px;
            background: linear-gradient(90deg, rgba(45,107,165,0.62) 0%, rgba(116,164,208,0.36) 100%);
        }
        .metric-card-title {
            font-size: 0.8rem;
            color: #4d6f91;
            font-weight: 680;
            margin-bottom: 0.24rem;
            letter-spacing: 0.15px;
        }
        .metric-card-value {
            font-size: 1.84rem;
            line-height: 1.08;
            color: #143a60;
            font-weight: 720;
        }
        .metric-card--alert .metric-card-title,
        .metric-card--alert .metric-card-value {
            color: #d92d20;
            font-weight: 800;
        }
        .analytics-hero {
            position: relative;
            overflow: hidden;
            border-radius: 20px;
            border: 1px solid rgba(45, 94, 146, 0.24);
            background:
                radial-gradient(860px 260px at 100% 0%, rgba(255, 255, 255, 0.22) 0%, rgba(255, 255, 255, 0) 60%),
                linear-gradient(135deg, #113558 0%, #1d5a90 38%, #3d7bb3 100%);
            padding: 22px 24px 18px 24px;
            color: #f3f8fe;
            box-shadow: 0 16px 36px rgba(16, 45, 79, 0.22);
            margin: 0.35rem 0 1rem 0;
        }
        .analytics-hero::after {
            content: "";
            position: absolute;
            right: -40px;
            top: -44px;
            width: 180px;
            height: 180px;
            border-radius: 999px;
            background: rgba(255, 255, 255, 0.10);
            filter: blur(6px);
        }
        .analytics-kicker {
            letter-spacing: 0.05rem;
            font-size: 0.82rem;
            color: rgba(240, 247, 253, 0.8);
            margin-bottom: 10px;
        }
        .analytics-title {
            font-size: 1.82rem;
            font-weight: 780;
            line-height: 1.15;
            margin-bottom: 8px;
        }
        .analytics-subtitle {
            font-size: 0.95rem;
            line-height: 1.55;
            color: rgba(244, 248, 252, 0.88);
            max-width: 980px;
            margin-bottom: 16px;
        }
        .analytics-chip-row {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 12px;
        }
        .analytics-chip {
            border-radius: 14px;
            border: 1px solid rgba(255, 255, 255, 0.18);
            background: linear-gradient(180deg, rgba(255,255,255,0.16) 0%, rgba(255,255,255,0.08) 100%);
            backdrop-filter: blur(6px);
            padding: 12px 14px;
            min-height: 86px;
        }
        .analytics-chip-label {
            font-size: 0.78rem;
            letter-spacing: 0.04rem;
            color: rgba(239, 246, 253, 0.82);
            margin-bottom: 4px;
        }
        .analytics-chip-value {
            font-size: 1.42rem;
            font-weight: 780;
            color: #ffffff;
            line-height: 1.1;
            margin-bottom: 6px;
        }
        .analytics-chip-note {
            font-size: 0.82rem;
            line-height: 1.45;
            color: rgba(239, 246, 253, 0.82);
        }
        .analytics-note {
            border-left: 4px solid #2e6ea9;
            border-radius: 0 14px 14px 0;
            background: linear-gradient(90deg, rgba(46,110,169,0.12) 0%, rgba(46,110,169,0.04) 100%);
            color: #173d63;
            padding: 10px 14px;
            margin: 0.35rem 0 0.8rem 0;
            font-size: 0.9rem;
            line-height: 1.55;
        }
        .analytics-note--soft {
            border-left-width: 3px;
            background: linear-gradient(90deg, rgba(89, 133, 176, 0.10) 0%, rgba(89, 133, 176, 0.03) 100%);
            color: #2a5279;
            margin-top: 0.2rem;
        }
        .analytics-section-map {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 12px;
            margin: 0.2rem 0 0.9rem 0;
        }
        .analytics-tab-top-gap {
            height: 24px;
        }
        .analytics-tab-block-gap {
            height: 20px;
        }
        .analytics-mini-divider {
            height: 1px;
            background: linear-gradient(90deg, rgba(167, 188, 210, 0.15) 0%, rgba(167, 188, 210, 0.75) 35%, rgba(167, 188, 210, 0.15) 100%);
            margin: 20px 0 18px 0;
        }
        .analytics-section-card {
            border-radius: 14px;
            border: 1px solid rgba(39, 92, 145, 0.12);
            background: linear-gradient(180deg, rgba(255,255,255,0.96) 0%, rgba(246,250,254,0.98) 100%);
            padding: 14px 15px 14px 15px;
            box-shadow:
                0 4px 10px rgba(16, 45, 79, 0.04),
                inset 0 1px 0 rgba(255,255,255,0.82);
            cursor: default;
        }
        .analytics-section-card-label {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            min-width: 46px;
            height: 24px;
            border-radius: 999px;
            border: 1px solid rgba(63, 110, 156, 0.16);
            background: linear-gradient(180deg, rgba(255,255,255,0.92) 0%, rgba(238,245,252,0.98) 100%);
            color: #456789;
            font-size: 0.75rem;
            font-weight: 780;
            letter-spacing: 0.02rem;
            margin-bottom: 8px;
        }
        .analytics-section-card-title {
            color: #163c61;
            font-size: 1.04rem;
            font-weight: 780;
            line-height: 1.25;
            margin-bottom: 6px;
        }
        .analytics-section-card-note {
            color: #4d6f91;
            font-size: 0.85rem;
            line-height: 1.56;
        }
        .analytics-subsection {
            display: flex;
            justify-content: space-between;
            align-items: end;
            gap: 16px;
            margin: 0.12rem 0 0.65rem 0;
            padding-bottom: 0.42rem;
            border-bottom: 1px solid rgba(157, 180, 205, 0.3);
        }
        .analytics-subsection-title {
            color: #163c61;
            font-size: 1.3rem;
            font-weight: 780;
            line-height: 1.2;
            margin-bottom: 5px;
        }
        .analytics-subsection-title .section-index-inline {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            min-width: 52px;
            height: 28px;
            margin-right: 10px;
            border-radius: 999px;
            border: 1px solid rgba(63, 110, 156, 0.18);
            background: linear-gradient(180deg, rgba(255,255,255,0.94) 0%, rgba(236,244,252,0.98) 100%);
            color: #3d6287;
            font-size: 0.84rem;
            font-weight: 800;
            letter-spacing: 0.02rem;
            vertical-align: middle;
        }
        .analytics-subsection-note {
            color: #5c7c9b;
            font-size: 0.85rem;
            line-height: 1.56;
            max-width: 760px;
        }
        .analytics-finale {
            border-radius: 18px;
            border: 1px solid rgba(42, 102, 161, 0.22);
            background:
                radial-gradient(420px 140px at 100% 0%, rgba(255,255,255,0.34) 0%, rgba(255,255,255,0) 70%),
                linear-gradient(135deg, #eff6ff 0%, #dfeeff 48%, #f8fbff 100%);
            padding: 16px 18px;
            margin: 0.4rem 0 0.95rem 0;
            box-shadow: 0 12px 24px rgba(20, 53, 89, 0.10);
        }
        .analytics-finale-title {
            color: #12385c;
            font-size: 1.02rem;
            font-weight: 760;
            margin-bottom: 6px;
        }
        .analytics-finale-body {
            color: #224b72;
            font-size: 0.92rem;
            line-height: 1.6;
        }
        .stage-hero {
            position: relative;
            overflow: hidden;
            border-radius: 20px;
            border: 1px solid rgba(45, 94, 146, 0.24);
            background:
                radial-gradient(640px 220px at 100% 0%, rgba(255,255,255,0.20) 0%, rgba(255,255,255,0) 60%),
                linear-gradient(135deg, #12375a 0%, #1f5f96 42%, #4f8dca 100%);
            padding: 20px 22px 16px 22px;
            color: #f4f8fe;
            box-shadow: 0 14px 30px rgba(16, 45, 79, 0.20);
            margin-bottom: 14px;
        }
        .stage-hero::after {
            content: "";
            position: absolute;
            right: -36px;
            top: -38px;
            width: 148px;
            height: 148px;
            border-radius: 999px;
            background: rgba(255, 255, 255, 0.10);
            filter: blur(6px);
        }
        .stage-kicker {
            letter-spacing: 0.05rem;
            font-size: 0.8rem;
            color: rgba(241, 247, 253, 0.8);
            margin-bottom: 8px;
        }
        .stage-title {
            font-size: 1.64rem;
            font-weight: 780;
            line-height: 1.14;
            margin-bottom: 7px;
        }
        .stage-subtitle {
            font-size: 0.93rem;
            line-height: 1.55;
            color: rgba(244, 248, 252, 0.86);
            max-width: 980px;
            margin-bottom: 14px;
        }
        .stage-meta-strip {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            margin: 0.1rem 0 0.65rem 0;
        }
        .stage-meta-pill {
            display: inline-flex;
            align-items: center;
            gap: 8px;
            border-radius: 8px;
            border: 1px solid rgba(45, 94, 146, 0.1);
            background: linear-gradient(180deg, rgba(255,255,255,0.9) 0%, rgba(246,249,253,0.98) 100%);
            color: #284b70;
            font-size: 0.84rem;
            line-height: 1.45;
            padding: 8px 12px;
            box-shadow: inset 0 1px 0 rgba(255,255,255,0.78);
            cursor: default;
        }
        .stage-meta-label {
            color: #5e7e9f;
            font-weight: 700;
        }
        .stage-inline-title {
            color: #35597d;
            font-size: 1.08rem;
            font-weight: 760;
            letter-spacing: 0.03rem;
            margin: 0.46rem 0 0.54rem 0;
            padding-left: 2px;
            border-left: 3px solid rgba(47, 110, 169, 0.42);
            padding-top: 2px;
            padding-bottom: 2px;
            padding-inline-start: 10px;
        }
        .stage-inline-title .stage-inline-index {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            min-width: 50px;
            height: 26px;
            margin-right: 10px;
            border-radius: 999px;
            border: 1px solid rgba(63, 110, 156, 0.16);
            background: linear-gradient(180deg, rgba(255,255,255,0.94) 0%, rgba(236,244,252,0.98) 100%);
            color: #3d6287;
            font-size: 0.82rem;
            font-weight: 800;
            letter-spacing: 0.02rem;
            vertical-align: middle;
        }
        .page-hero {
            position: relative;
            overflow: hidden;
            border-radius: 24px;
            border: 1px solid rgba(93, 99, 112, 0.16);
            background:
                radial-gradient(880px 280px at 100% 0%, rgba(255,255,255,0.82) 0%, rgba(255,255,255,0) 62%),
                linear-gradient(180deg, #fbfaf7 0%, #f2efe8 100%);
            padding: 24px 26px 24px 26px;
            margin: 0.2rem 0 1.1rem 0;
            color: #1d3147;
            box-shadow: 0 14px 28px rgba(44, 52, 64, 0.10);
        }
        .page-hero::after {
            content: "";
            position: absolute;
            right: -56px;
            top: -60px;
            width: 220px;
            height: 220px;
            border-radius: 999px;
            background: rgba(170, 182, 196, 0.16);
            filter: blur(8px);
        }
        .page-hero-kicker {
            letter-spacing: 0.05rem;
            font-size: 0.84rem;
            color: #6a7280;
            margin-bottom: 8px;
        }
        .page-hero-title {
            font-size: 1.88rem;
            line-height: 1.14;
            font-weight: 800;
            margin-bottom: 8px;
            color: #1b3046;
        }
        .page-hero-subtitle {
            font-size: 0.95rem;
            line-height: 1.58;
            color: #536275;
            max-width: 980px;
            margin-bottom: 20px;
        }
        .page-hero-flow {
            display: grid;
            grid-template-columns: minmax(0, 1fr) 34px minmax(0, 1fr) 34px minmax(0, 1fr) 34px minmax(0, 1fr);
            gap: 10px;
            align-items: stretch;
        }
        .page-hero-step {
            position: relative;
            border-radius: 16px;
            border: 1px solid rgba(120, 134, 150, 0.18);
            background:
                radial-gradient(180px 70px at 100% 0%, rgba(226, 234, 242, 0.46) 0%, rgba(226, 234, 242, 0) 70%),
                linear-gradient(180deg, rgba(255,255,255,0.96) 0%, rgba(247,245,240,0.98) 100%);
            padding: 16px 16px 16px 16px;
            min-height: 170px;
            z-index: 1;
            box-shadow:
                0 10px 20px rgba(39, 48, 60, 0.06),
                inset 0 1px 0 rgba(255,255,255,0.84);
        }
        .page-hero-connector {
            display: flex;
            align-items: center;
            justify-content: center;
            min-height: 170px;
        }
        .page-hero-connector-line {
            position: relative;
            width: 100%;
            height: 2px;
            border-radius: 999px;
            background: linear-gradient(90deg, rgba(150, 164, 182, 0.18) 0%, rgba(118, 135, 155, 0.68) 45%, rgba(118, 135, 155, 0.68) 100%);
        }
        .page-hero-connector-line::after {
            content: "";
            position: absolute;
            right: -1px;
            top: 50%;
            width: 10px;
            height: 10px;
            border-top: 2px solid #76879b;
            border-right: 2px solid #76879b;
            transform: translateY(-50%) rotate(45deg);
        }
        .page-hero-step-head {
            display: flex;
            align-items: center;
            gap: 12px;
            margin-bottom: 13px;
        }
        .page-hero-step-no {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 42px;
            height: 42px;
            border-radius: 999px;
            background: linear-gradient(180deg, #c85b53 0%, #a63f38 100%);
            color: #ffffff;
            font-size: 0.98rem;
            font-weight: 800;
            flex: 0 0 42px;
            box-shadow: 0 8px 16px rgba(166, 63, 56, 0.18);
        }
        .page-hero-step-headline {
            font-size: 1.38rem;
            font-weight: 780;
            color: #1c3249;
            line-height: 1.22;
            margin-top: 1px;
        }
        .page-hero-step-body {
            font-size: 0.99rem;
            line-height: 1.58;
            color: #56667a;
            max-width: 26ch;
        }
        @media (max-width: 1100px) {
            .page-hero-flow {
                grid-template-columns: repeat(2, minmax(0, 1fr));
            }
            .page-hero-connector {
                display: none;
            }
        }
        @media (max-width: 720px) {
            .page-hero-flow {
                grid-template-columns: 1fr;
            }
        }
        .legend-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 10px;
            margin: 0.2rem 0 0.85rem 0;
        }
        .legend-chip {
            display: flex;
            align-items: center;
            gap: 10px;
            border-radius: 12px;
            border: 1px solid rgba(43, 94, 145, 0.12);
            background: linear-gradient(180deg, rgba(255,255,255,0.88) 0%, rgba(242,247,253,0.96) 100%);
            padding: 10px 12px;
            color: #1f4468;
            font-size: 0.86rem;
            line-height: 1.45;
            box-shadow: inset 0 1px 0 rgba(255,255,255,0.72);
            cursor: default;
        }
        .legend-swatch {
            width: 16px;
            height: 16px;
            border-radius: 5px;
            border: 1px solid rgba(16, 43, 71, 0.16);
            flex: 0 0 16px;
        }
        .switcher-panel-title {
            color: #557595;
            font-size: 0.86rem;
            font-weight: 700;
            letter-spacing: 0.04rem;
            margin: 0.15rem 0 0.48rem 0;
        }
        .info-panel {
            border-radius: 14px;
            border: 1px solid rgba(39, 92, 145, 0.12);
            background: linear-gradient(180deg, rgba(255,255,255,0.96) 0%, rgba(244,248,253,0.98) 100%);
            padding: 10px 12px;
            box-shadow:
                0 6px 14px rgba(16, 45, 79, 0.04),
                inset 0 1px 0 rgba(255,255,255,0.82);
        }
        .info-panel-row {
            display: grid;
            grid-template-columns: 88px minmax(0, 1fr);
            gap: 10px;
            align-items: start;
            padding: 6px 0;
        }
        .info-panel-row + .info-panel-row {
            border-top: 1px solid rgba(189, 209, 229, 0.42);
        }
        .info-panel-label {
            color: #5a7a99;
            font-size: 0.82rem;
            font-weight: 700;
            line-height: 1.5;
        }
        .info-panel-value {
            color: #1c446a;
            font-size: 0.88rem;
            font-weight: 600;
            line-height: 1.58;
            word-break: break-word;
        }
        [data-testid="stDataFrame"] {
            border-radius: 16px;
            border: 1px solid rgba(39, 92, 145, 0.14);
            background:
                radial-gradient(240px 90px at 100% 0%, rgba(255,255,255,0.38) 0%, rgba(255,255,255,0) 70%),
                linear-gradient(180deg, rgba(255,255,255,0.92) 0%, rgba(243,248,253,0.98) 100%);
            box-shadow:
                0 8px 18px rgba(16, 45, 79, 0.06),
                inset 0 1px 0 rgba(255,255,255,0.8);
            padding: 4px;
        }
        [data-testid="stDataFrame"] [role="columnheader"] {
            background: linear-gradient(180deg, #eef5fd 0%, #e0ecf8 100%) !important;
            color: #24486d !important;
            font-weight: 760 !important;
            border-bottom: 1px solid rgba(169, 194, 221, 0.72) !important;
        }
        [data-testid="stDataFrame"] [role="gridcell"] {
            color: #23476d !important;
            background: rgba(255,255,255,0.76) !important;
        }
        [data-testid="stAlert"] {
            border-radius: 14px;
            border-width: 1px;
            box-shadow: 0 8px 16px rgba(16, 45, 79, 0.05);
        }
        [data-testid="stMarkdownContainer"] p {
            line-height: 1.62;
        }
        .toolbar-side-shell {
            border-radius: 18px;
            border: 1px solid rgba(39, 92, 145, 0.14);
            background:
                radial-gradient(260px 100px at 100% 0%, rgba(255,255,255,0.34) 0%, rgba(255,255,255,0) 70%),
                linear-gradient(180deg, rgba(255,255,255,0.84) 0%, rgba(237,244,252,0.96) 100%);
            padding: 16px 16px 14px 16px;
            margin-bottom: 22px;
            box-shadow:
                0 10px 24px rgba(16, 45, 79, 0.06),
                inset 0 1px 0 rgba(255,255,255,0.76);
        }
        .toolbar-side-title {
            color: #163c61;
            font-size: 1.2rem;
            font-weight: 780;
            line-height: 1.2;
            margin-bottom: 8px;
        }
        .toolbar-side-note {
            color: #59799a;
            font-size: 0.86rem;
            line-height: 1.55;
            margin-bottom: 16px;
        }
        .toolbar-side-count {
            border-radius: 10px;
            border: 1px solid rgba(45, 94, 146, 0.12);
            background: linear-gradient(180deg, rgba(255,255,255,0.94) 0%, rgba(244,248,253,0.98) 100%);
            color: #284b70;
            font-size: 0.9rem;
            font-weight: 700;
            padding: 14px 16px;
            margin-top: 10px;
            margin-bottom: 18px;
        }
        .toolbar-side-divider {
            height: 1px;
            background: linear-gradient(90deg, rgba(181, 201, 223, 0.16) 0%, rgba(181, 201, 223, 0.75) 35%, rgba(181, 201, 223, 0.16) 100%);
            margin: 14px 0 14px 0;
        }
        .toolbar-row-title {
            color: #365a7f;
            font-size: 1.42rem;
            font-weight: 780;
            line-height: 1.28;
            margin: 0.22rem 0 0.8rem 0;
        }
        .toolbar-row-spacer {
            height: 24px;
        }
        .toolbar-side-footer {
            margin-top: 36px;
        }
        .section-gap-sm {
            height: 14px;
        }
        .selection-summary {
            display: inline-flex;
            align-items: center;
            gap: 8px;
            border-radius: 10px;
            border: 1px solid rgba(45, 94, 146, 0.12);
            background: linear-gradient(180deg, rgba(255,255,255,0.94) 0%, rgba(244,248,253,0.98) 100%);
            color: #284b70;
            font-size: 0.86rem;
            font-weight: 700;
            line-height: 1.5;
            padding: 8px 12px;
            margin-bottom: 10px;
            box-shadow: inset 0 1px 0 rgba(255,255,255,0.82);
        }
        .selection-summary strong {
            color: #153a5f;
            font-size: 1rem;
            font-weight: 800;
        }
        .result-shell {
            border-radius: 20px;
            border: 1px solid rgba(31, 78, 122, 0.18);
            background:
                radial-gradient(260px 100px at 100% 0%, rgba(255,255,255,0.34) 0%, rgba(255,255,255,0) 70%),
                linear-gradient(180deg, rgba(255,255,255,0.82) 0%, rgba(236,244,252,0.96) 100%);
            padding: 16px 16px 14px 16px;
            box-shadow:
                0 14px 28px rgba(18, 53, 89, 0.08),
                inset 0 1px 0 rgba(255,255,255,0.78);
        }
        .result-hero {
            border-radius: 16px;
            border: 1px solid rgba(34, 92, 143, 0.18);
            background:
                radial-gradient(180px 80px at 100% 0%, rgba(255,255,255,0.24) 0%, rgba(255,255,255,0) 65%),
                linear-gradient(135deg, #18466f 0%, #2f6ea9 52%, #5b96ce 100%);
            padding: 14px 14px 12px 14px;
            margin-bottom: 12px;
            box-shadow: 0 12px 24px rgba(18, 53, 89, 0.18);
        }
        .result-hero-title {
            color: rgba(240, 247, 253, 0.82);
            font-size: 0.78rem;
            text-transform: uppercase;
            letter-spacing: 0.14rem;
            margin-bottom: 6px;
        }
        .result-hero-value {
            color: #ffffff;
            font-size: 1.56rem;
            font-weight: 800;
            line-height: 1.08;
            margin-bottom: 5px;
        }
        .result-hero-note {
            color: rgba(240, 247, 253, 0.84);
            font-size: 0.84rem;
            line-height: 1.5;
        }
        .analytics-progress {
            margin: 0.15rem 0 0.55rem 0;
        }
        .playback-head {
            display: flex;
            justify-content: space-between;
            align-items: center;
            gap: 12px;
            margin: 0.1rem 0 0.35rem 0;
        }
        .playback-head-title {
            color: #35597d;
            font-size: 0.9rem;
            font-weight: 760;
            letter-spacing: 0.03rem;
        }
        .playback-head-status {
            border-radius: 999px;
            border: 1px solid rgba(45, 94, 146, 0.12);
            background: linear-gradient(180deg, rgba(255,255,255,0.92) 0%, rgba(244,248,253,0.98) 100%);
            color: #527496;
            font-size: 0.78rem;
            font-weight: 700;
            padding: 6px 10px;
            line-height: 1;
        }
        .major-section-divider {
            height: 1px;
            margin: 1.2rem 0 1.45rem 0;
            background: linear-gradient(90deg, rgba(160, 176, 194, 0.12) 0%, rgba(160, 176, 194, 0.82) 28%, rgba(160, 176, 194, 0.82) 72%, rgba(160, 176, 194, 0.12) 100%);
        }
        .analytics-progress-meta {
            display: flex;
            justify-content: space-between;
            align-items: center;
            color: #426788;
            font-size: 0.8rem;
            margin-bottom: 6px;
        }
        .analytics-progress-track {
            position: relative;
            height: 7px;
            border-radius: 999px;
            background: linear-gradient(90deg, rgba(191, 210, 229, 0.86) 0%, rgba(222, 234, 246, 0.96) 100%);
            overflow: hidden;
            box-shadow: inset 0 1px 2px rgba(14, 49, 84, 0.10);
        }
        .analytics-progress-fill {
            height: 100%;
            border-radius: 999px;
            background: linear-gradient(90deg, #2d6ba5 0%, #5b96ce 55%, #86b8e5 100%);
            box-shadow: 0 0 14px rgba(61, 121, 180, 0.28);
            transition: width 0.22s ease;
        }
        [data-testid="stCaptionContainer"] {
            color: #446587;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_selected_ids_panel(selected_ids: List[int]) -> None:
    import streamlit as st

    if not selected_ids:
        st.markdown('<div class="selected-empty">未选择点位</div>', unsafe_allow_html=True)
        return

    chips = "".join(f'<span class="selected-chip">{point_id}</span>' for point_id in selected_ids)
    st.markdown(
        f'<div class="selected-box"><div class="selected-grid">{chips}</div></div>',
        unsafe_allow_html=True,
    )


def render_toolbar_selected_ids_panel(selected_ids: List[int]) -> None:
    import streamlit as st

    if not selected_ids:
        st.markdown(
            '<div class="selected-empty toolbar-selected-box">未选择点位</div>',
            unsafe_allow_html=True,
        )
        return

    chips = "".join(f'<span class="selected-chip">{point_id}</span>' for point_id in selected_ids)
    st.markdown(
        (
            '<div class="selected-box toolbar-selected-box">'
            f'<div class="selected-grid toolbar-selected-grid">{chips}</div>'
            "</div>"
        ),
        unsafe_allow_html=True,
    )


def render_detection_details_panel(
    polluted_ids: List[int],
    glove_names: List[str],
    report_status_items: Optional[List[Tuple[str, str]]] = None,
) -> None:
    import streamlit as st

    def chunk_join(items: List[str], chunk_size: int) -> str:
        if not items:
            return "未识别到"
        lines = ["、".join(items[i : i + chunk_size]) for i in range(0, len(items), chunk_size)]
        return "<br/>".join(html.escape(line) for line in lines)

    floor_items = [str(point_id) for point_id in polluted_ids]
    glove_items = [name for name in glove_names]
    mapping_items = [
        f"ID {point_id} -> 行{id_to_row_col(point_id)[0]} 列{id_to_row_col(point_id)[1]}"
        for point_id in polluted_ids
    ]

    floor_text = chunk_join(floor_items, 8)
    glove_text = chunk_join(glove_items, 4)
    mapping_text = chunk_join(mapping_items, 2)
    status_markup = ""
    if report_status_items:
        rows = "".join(
            (
                '<div class="detail-info-row">'
                f'<div class="detail-info-label">{html.escape(label)}</div>'
                f'<div class="detail-info-value">{html.escape(value)}</div>'
                "</div>"
            )
            for label, value in report_status_items
        )
        status_markup = (
            '<div class="detail-row">'
            '<div class="detail-title">3.2.4 当前报告状态 <span class="detail-anchor">•</span></div>'
            '<div class="detail-item">'
            f'<div class="detail-body detail-info-rows">{rows}</div>'
            "</div>"
            "</div>"
        )

    st.markdown(
        f"""
        <div class="detail-stack">
            <div class="detail-row">
                <div class="detail-title">3.2.1 当前污染点位 ID <span class="detail-anchor">•</span></div>
                <div class="detail-item">
                    <div class="detail-body">{floor_text}</div>
                </div>
            </div>
            <div class="detail-row">
                <div class="detail-title">3.2.2 超标手套孔名称 <span class="detail-anchor">•</span></div>
                <div class="detail-item">
                    <div class="detail-body">{glove_text}</div>
                </div>
            </div>
            <div class="detail-row">
                <div class="detail-title">3.2.3 点位对应行列 <span class="detail-anchor">•</span></div>
                <div class="detail-item">
                    <div class="detail-body">{mapping_text}</div>
                </div>
            </div>
            {status_markup}
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_metric_card(label: str, value: int, alert: bool = False) -> None:
    import streamlit as st

    alert_class = " metric-card--alert" if alert else ""
    st.markdown(
        f"""
        <div class="metric-card{alert_class}">
            <div class="metric-card-title">{label}</div>
            <div class="metric-card-value">{value}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_export_button(label: str, rows: List[Dict[str, Any]], file_name: str, key: str) -> None:
    import streamlit as st

    st.download_button(
        label=label,
        data=rows_to_csv_bytes(rows),
        file_name=file_name,
        mime="text/csv",
        key=key,
        use_container_width=False,
    )


def render_stage_panel_start(
    kicker: str,
    title: str,
    subtitle: str,
    chips: Optional[List[Tuple[str, str, str]]] = None,
) -> None:
    import streamlit as st

    chip_markup = ""
    if chips:
        chip_markup = '<div class="analytics-chip-row">' + "".join(
            (
                '<div class="analytics-chip">'
                f'<div class="analytics-chip-label">{html.escape(label)}</div>'
                f'<div class="analytics-chip-value">{html.escape(value)}</div>'
                f'<div class="analytics-chip-note">{html.escape(note)}</div>'
                "</div>"
            )
            for label, value, note in chips
        ) + "</div>"

    st.markdown(
        f"""
        <div class="stage-hero">
            <div class="stage-kicker">{html.escape(kicker)}</div>
            <div class="stage-title">{html.escape(title)}</div>
            <div class="stage-subtitle">{html.escape(subtitle)}</div>
            {chip_markup}
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_stage_panel_end() -> None:
    return None


def render_stage_subpanel_start(title: str, index: str = "") -> None:
    import streamlit as st

    index_markup = f'<span class="stage-inline-index">{html.escape(index)}</span>' if index else ""
    st.markdown(
        f'<div class="stage-inline-title">{index_markup}{html.escape(title)}</div>',
        unsafe_allow_html=True,
    )


def render_stage_subpanel_end() -> None:
    return None


def render_stage_meta_strip(items: List[Tuple[str, str]]) -> None:
    import streamlit as st

    markup = "".join(
        (
            '<div class="stage-meta-pill">'
            f'<span class="stage-meta-label">{html.escape(label)}</span>'
            f"<span>{html.escape(value)}</span>"
            "</div>"
        )
        for label, value in items
    )
    st.markdown(f'<div class="stage-meta-strip">{markup}</div>', unsafe_allow_html=True)


def render_info_panel(items: List[Tuple[str, str]]) -> None:
    import streamlit as st

    markup = "".join(
        (
            '<div class="info-panel-row">'
            f'<div class="info-panel-label">{html.escape(label)}</div>'
            f'<div class="info-panel-value">{html.escape(value)}</div>'
            "</div>"
        )
        for label, value in items
    )
    st.markdown(f'<div class="info-panel">{markup}</div>', unsafe_allow_html=True)


def render_page_hero(steps: List[Tuple[str, str, str, str]]) -> None:
    import streamlit as st

    flow_parts: List[str] = []
    for idx, (step_no, title, task, gain) in enumerate(steps):
        flow_parts.append(
            '<div class="page-hero-step">'
            '<div class="page-hero-step-head">'
            f'<div class="page-hero-step-no">{html.escape(step_no)}</div>'
            f'<div class="page-hero-step-headline">{html.escape(title)}</div>'
            "</div>"
            f'<div class="page-hero-step-body">{html.escape(task)} {html.escape(gain)}</div>'
            "</div>"
        )
        if idx < len(steps) - 1:
            flow_parts.append('<div class="page-hero-connector"><div class="page-hero-connector-line"></div></div>')
    flow_markup = "".join(flow_parts)
    st.markdown(
        f"""
        <div class="page-hero">
            <div class="page-hero-kicker">功能说明</div>
            <div class="page-hero-title">辐射超标地面点位与手套孔智能研判平台</div>
            <div class="page-hero-subtitle">
                平台提供报告接入、区域修正、结果展示和趋势分析四类核心能力。
                可按照以下顺序完成数据导入、网格修正、结果核查与趋势研判。
            </div>
            <div class="page-hero-flow">{flow_markup}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_switcher_panel(
    title: str, switch_id: str, options: List[Tuple[str, str]], active_key: str
) -> None:
    import streamlit as st

    st.markdown(f'<div class="switcher-panel-title">{html.escape(title)}</div>', unsafe_allow_html=True)
    render_report_switcher("", switch_id, options, active_key)


def render_result_hero(polluted_count: int, glove_count: int, selected_count: int) -> None:
    import streamlit as st

    st.markdown(
        f"""
        <div class="result-hero">
            <div class="result-hero-title">结果总览</div>
            <div class="result-hero-value">{polluted_count} 个污染点位</div>
            <div class="result-hero-note">
                当前识别到 {glove_count} 个超标手套孔，暂存选区 {selected_count} 个点位。
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_terminology_legend() -> None:
    import streamlit as st

    st.markdown(
        """
        <div class="analytics-note">
            术语统一说明：污染点位 = 当前经报告识别或人工修正后确认的污染格；无污染点位 = 可达且未标记为污染的格；
            障碍物点位 = 不参与作业路径的阻挡区域；待测区域点位 = 暂不纳入污染/无污染判断、等待复核的区域。
        </div>
        <div class="legend-grid">
            <div class="legend-chip"><span class="legend-swatch" style="background:#ff4d4f;"></span>污染点位</div>
            <div class="legend-chip"><span class="legend-swatch" style="background:#2e6ea9;"></span>当前框选</div>
            <div class="legend-chip"><span class="legend-swatch" style="background:#d2f4d2;"></span>无污染点位</div>
            <div class="legend-chip"><span class="legend-swatch" style="background:#8d9096;"></span>障碍物点位</div>
            <div class="legend-chip"><span class="legend-swatch" style="background:#fff7cc;"></span>待测区域点位</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_analytics_hero(
    latest: Dict[str, Any],
    forecast_count: int,
    stable_hotspots: int,
    top_risk: Optional[Dict[str, Any]],
    top_source: Optional[Dict[str, Any]],
) -> None:
    import streamlit as st

    latest_count = int(latest["polluted_count"])
    delta = forecast_count - latest_count
    trend_text = "扩张" if delta > 0 else "收敛" if delta < 0 else "持平"
    top_risk_text = (
        f"ID {top_risk['point_id']} / 风险 {top_risk['score'] * 100:.1f}"
        if top_risk
        else "暂无高风险点"
    )
    top_source_text = (
        f"ID {top_source['point_id']} / 评分 {top_source['score'] * 100:.1f}"
        if top_source
        else "暂无疑似源点"
    )
    source_note = top_source["explanation"] if top_source else "当前历史数据不足以判定明显源点。"

    st.markdown(
        f"""
        <div class="analytics-hero">
            <div class="analytics-kicker">分析总览</div>
            <div class="analytics-title">趋势预测、统计计算与污染溯源</div>
            <div class="analytics-subtitle">
                系统根据多日报告的污染点位分布，综合评估趋势走向、扩散风险与疑似源头。
                当前判断为 <strong>{trend_text}</strong>，下一期预测较最新一期 <strong>{delta:+d}</strong> 个点位。
            </div>
            <div class="analytics-chip-row">
                <div class="analytics-chip">
                    <div class="analytics-chip-label">下一期趋势</div>
                    <div class="analytics-chip-value">{forecast_count}</div>
                    <div class="analytics-chip-note">相较当前 {latest_count} 个污染点位，趋势判断为 {trend_text}</div>
                </div>
                <div class="analytics-chip">
                    <div class="analytics-chip-label">持续热点</div>
                    <div class="analytics-chip-value">{stable_hotspots}</div>
                    <div class="analytics-chip-note">连续出现两期及以上的稳定热点数量</div>
                </div>
                <div class="analytics-chip">
                    <div class="analytics-chip-label">最高风险点</div>
                    <div class="analytics-chip-value">{top_risk_text}</div>
                    <div class="analytics-chip-note">优先复核扩散边界与邻域联动</div>
                </div>
                <div class="analytics-chip">
                    <div class="analytics-chip-label">疑似源点</div>
                    <div class="analytics-chip-value">{top_source_text}</div>
                    <div class="analytics-chip-note">{html.escape(source_note)}</div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_playback_completion_card(title: str, body: str) -> None:
    import streamlit as st

    st.markdown(
        f"""
        <div class="analytics-finale">
            <div class="analytics-finale-title">{html.escape(title)}</div>
            <div class="analytics-finale-body">{html.escape(body)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_playback_progress(current_index: int, total_count: int, current_label: str) -> None:
    import streamlit as st

    safe_total = max(total_count, 1)
    progress_ratio = min(max((current_index + 1) / safe_total, 0.0), 1.0)
    st.markdown(
        f"""
        <div class="analytics-progress">
            <div class="analytics-progress-meta">
                <span>回放进度</span>
                <span>第 {current_index + 1} / {safe_total} 期 · {html.escape(current_label)}</span>
            </div>
            <div class="analytics-progress-track">
                <div class="analytics-progress-fill" style="width: {progress_ratio * 100:.1f}%;"></div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_section_path_map(items: List[Tuple[str, str, str]]) -> None:
    import streamlit as st

    cards_markup = "".join(
        (
            '<div class="analytics-section-card">'
            f'<div class="analytics-section-card-label">{html.escape(label)}</div>'
            f'<div class="analytics-section-card-title">{html.escape(title)}</div>'
            f'<div class="analytics-section-card-note">{html.escape(note)}</div>'
            "</div>"
        )
        for label, title, note in items
    )
    st.markdown(f'<div class="analytics-section-map">{cards_markup}</div>', unsafe_allow_html=True)


def render_analytics_section_map() -> None:
    render_section_path_map(
        [
            ("5.1", "趋势总览", "先看总量变化、污染占比、重心迁移和结构集中度。"),
            ("5.2", "热点画像", "再看哪些点位反复出现、持续时间更久、历史频率更高。"),
            ("5.3", "趋势预测", "然后根据历史频率、邻域支撑和回放快照推演下一期风险。"),
            ("5.4", "污染溯源", "最后综合首次出现、持续性和扩散链路给出优先排查候选点。"),
        ]
    )


def render_daily_result_section_map() -> None:
    render_section_path_map(
        [
            ("3.1", "统计概览", "先看当前日期下的污染点位、手套孔、障碍物和无污染点位总数。"),
            ("3.2", "结果明细", "再看污染点位 ID、手套孔名称以及点位对应的行列位置明细。"),
        ]
    )


def render_section_subsection(title: str, note: str, tag: str) -> None:
    import streamlit as st

    st.markdown(
        f"""
        <div class="analytics-subsection">
            <div>
                <div class="analytics-subsection-title"><span class="section-index-inline">{html.escape(tag)}</span>{html.escape(title)}</div>
                <div class="analytics-subsection-note">{html.escape(note)}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_analytics_subsection(title: str, note: str, tag: str) -> None:
    render_section_subsection(title, note, tag)


def render_note_bar(text: str, soft: bool = False) -> None:
    import streamlit as st

    extra_class = " analytics-note--soft" if soft else ""
    st.markdown(
        f'<div class="analytics-note{extra_class}">{html.escape(text)}</div>',
        unsafe_allow_html=True,
    )


def render_analysis_methodology(
    title: str,
    paragraphs: List[str],
    formulas: Optional[List[str]] = None,
    formula_explanations: Optional[List[List[str]]] = None,
) -> None:
    import streamlit as st

    st.markdown(f"#### {title}")
    for paragraph in paragraphs:
        render_note_bar(paragraph, soft=True)
    if formulas:
        for index, formula in enumerate(formulas):
            st.latex(formula)
            if formula_explanations and index < len(formula_explanations):
                for explanation in formula_explanations[index]:
                    render_note_bar(explanation, soft=True)


def render_analytics_mini_divider() -> None:
    import streamlit as st

    st.markdown('<div class="analytics-mini-divider"></div>', unsafe_allow_html=True)


def render_selection_summary(count: int) -> None:
    import streamlit as st

    st.markdown(
        f'<div class="selection-summary">当前暂存选区 <strong>{count}</strong> 个点位</div>',
        unsafe_allow_html=True,
    )


def render_major_section_divider() -> None:
    import streamlit as st

    st.markdown('<div class="major-section-divider"></div>', unsafe_allow_html=True)


def render_analysis_config_panel() -> Dict[str, float]:
    import streamlit as st

    defaults = default_analysis_config()
    with st.expander("分析参数与假设", expanded=False):
        render_note_bar("这里可以直接调整趋势预测、风险预测、溯源评分的权重。滑块改动后图表会自动重算，方便根据现场经验校准模型。")
        forecast_tab, risk_tab, source_tab = st.tabs(["趋势预测", "风险预测", "污染溯源"])

        with forecast_tab:
            c1, c2 = st.columns(2)
            with c1:
                forecast_window = st.slider("预测窗口期数", 2, 8, int(defaults["forecast_window"]))
                forecast_smooth_weight = st.slider(
                    "平滑趋势权重", 0.0, 1.0, float(defaults["forecast_smooth_weight"]), 0.02
                )
            with c2:
                forecast_drift_scale = st.slider(
                    "漂移放大量级", 0.0, 2.0, float(defaults["forecast_drift_scale"]), 0.05
                )
                forecast_trend_weight = st.slider(
                    "近期漂移权重", 0.0, 1.0, float(defaults["forecast_trend_weight"]), 0.02
                )

        with risk_tab:
            c1, c2 = st.columns(2)
            with c1:
                risk_weight_frequency = st.slider("历史频率权重", 0.0, 1.0, float(defaults["risk_weight_frequency"]), 0.02)
                risk_weight_recent = st.slider("近三期活跃权重", 0.0, 1.0, float(defaults["risk_weight_recent"]), 0.02)
                risk_weight_streak = st.slider("当前连续性权重", 0.0, 1.0, float(defaults["risk_weight_streak"]), 0.02)
                risk_threshold = st.slider("风险入围阈值", 0.0, 0.6, float(defaults["risk_threshold"]), 0.01)
            with c2:
                risk_weight_neighbor = st.slider("邻域支撑权重", 0.0, 1.0, float(defaults["risk_weight_neighbor"]), 0.02)
                risk_weight_regional = st.slider("区域扩散权重", 0.0, 1.0, float(defaults["risk_weight_regional"]), 0.02)
                risk_bonus_current = st.slider("当前污染加分", 0.0, 0.3, float(defaults["risk_bonus_current"]), 0.01)
                risk_bonus_spread = st.slider("边界扩散加分", 0.0, 0.3, float(defaults["risk_bonus_spread"]), 0.01)

        with source_tab:
            c1, c2 = st.columns(2)
            with c1:
                source_weight_first = st.slider("首次出现权重", 0.0, 1.0, float(defaults["source_weight_first"]), 0.02)
                source_weight_persistence = st.slider(
                    "持续污染权重", 0.0, 1.0, float(defaults["source_weight_persistence"]), 0.02
                )
                source_weight_spread = st.slider("周边扩散权重", 0.0, 1.0, float(defaults["source_weight_spread"]), 0.02)
                source_spread_radius = st.slider("扩散搜索半径", 1, 6, int(defaults["source_spread_radius"]))
            with c2:
                source_weight_frequency = st.slider(
                    "历史频率权重", 0.0, 1.0, float(defaults["source_weight_frequency"]), 0.02
                )
                source_weight_centroid = st.slider(
                    "热点中心权重", 0.0, 1.0, float(defaults["source_weight_centroid"]), 0.02
                )
                source_spread_norm = st.slider("扩散归一化阈值", 2, 20, int(defaults["source_spread_norm"]))
                source_centroid_norm = st.slider("中心距离归一化", 6, 30, int(defaults["source_centroid_norm"]))

    return {
        "forecast_window": float(forecast_window),
        "forecast_smooth_weight": float(forecast_smooth_weight),
        "forecast_trend_weight": float(forecast_trend_weight),
        "forecast_drift_scale": float(forecast_drift_scale),
        "risk_weight_frequency": float(risk_weight_frequency),
        "risk_weight_recent": float(risk_weight_recent),
        "risk_weight_streak": float(risk_weight_streak),
        "risk_weight_neighbor": float(risk_weight_neighbor),
        "risk_weight_regional": float(risk_weight_regional),
        "risk_bonus_current": float(risk_bonus_current),
        "risk_bonus_spread": float(risk_bonus_spread),
        "risk_threshold": float(risk_threshold),
        "source_weight_first": float(source_weight_first),
        "source_weight_persistence": float(source_weight_persistence),
        "source_weight_spread": float(source_weight_spread),
        "source_weight_frequency": float(source_weight_frequency),
        "source_weight_centroid": float(source_weight_centroid),
        "source_spread_radius": float(source_spread_radius),
        "source_spread_norm": float(source_spread_norm),
        "source_centroid_norm": float(source_centroid_norm),
    }


def render_playback_selector(history: List[Dict[str, Any]], key: str, title: str) -> int:
    import streamlit as st

    labels = build_playback_labels(history)
    selected_label = st.select_slider(
        title,
        options=labels,
        value=labels[-1],
        key=key,
    )
    return labels.index(selected_label)


def _format_axis_date_label(value: Any) -> str:
    text = str(value or "").strip()
    match = re.match(r"^(\d{4})(\d{2})(\d{2})(.*)$", text)
    if not match:
        return text
    suffix = (match.group(4) or "").strip()
    base = f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
    return f"{base}{suffix}"


def _build_history_x_labels(history: List[Dict[str, Any]]) -> List[str]:
    labels: List[str] = []
    for item in history:
        switch_label = str(item.get("switch_label") or item.get("date") or "未识别日期")
        labels.append(_format_axis_date_label(switch_label))
    return labels


def _centered_chart_title(text: str, y: float = 0.97) -> Dict[str, Any]:
    return {
        "text": text,
        "x": 0.5,
        "xanchor": "center",
        "y": y,
        "yanchor": "top",
    }


def build_pollution_count_figure(history: List[Dict[str, Any]], forecast_count: int):
    import plotly.graph_objects as go

    labels = _build_history_x_labels(history)
    counts = [int(item["polluted_count"]) for item in history]
    prediction_label = "预测下一期"

    fig = go.Figure()
    fig.add_trace(
        go.Scatter(
            x=labels,
            y=counts,
            name="实际污染点位",
            mode="lines+markers",
            line={"color": "#e14b43", "width": 3},
            marker={"size": 8, "color": "#ffffff", "line": {"width": 2, "color": "#e14b43"}},
        )
    )
    if labels:
        fig.add_trace(
            go.Scatter(
                x=[labels[-1], prediction_label],
                y=[counts[-1], int(forecast_count)],
                name="下一期预测",
                mode="lines+markers",
                line={"color": "#2e6ea9", "width": 2.6, "dash": "dot"},
                marker={"size": 8, "color": "#ffffff", "line": {"width": 2, "color": "#2e6ea9"}},
            )
        )

    fig.update_layout(
        title=_centered_chart_title("污染点位趋势与下一期预测", y=0.975),
        height=320,
        margin={"l": 28, "r": 28, "t": 88, "b": 34},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.92)",
        legend={"orientation": "h", "yanchor": "bottom", "y": 1.11, "x": 0.5, "xanchor": "center"},
        hovermode="x unified",
    )
    fig.update_xaxes(type="category", tickangle=-22, categoryorder="array", categoryarray=labels + [prediction_label])
    fig.update_yaxes(title_text="污染点位数", rangemode="tozero", zerolinecolor="#dce8f4")
    return fig


def build_added_points_figure(history: List[Dict[str, Any]]):
    import plotly.graph_objects as go

    labels = _build_history_x_labels(history)
    added = [len(item["added_ids"]) for item in history]
    fig = go.Figure()
    fig.add_trace(
        go.Bar(
            x=labels,
            y=added,
            name="新增点位",
            marker_color="#50b26b",
            opacity=0.86,
        )
    )
    fig.update_layout(
        title=_centered_chart_title("新增点位变化"),
        height=300,
        margin={"l": 28, "r": 28, "t": 56, "b": 34},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.92)",
        hovermode="x unified",
    )
    fig.update_xaxes(type="category", tickangle=-22, categoryorder="array", categoryarray=labels)
    fig.update_yaxes(title_text="新增点位数", rangemode="tozero")
    return fig


def build_removed_points_figure(history: List[Dict[str, Any]]):
    import plotly.graph_objects as go

    labels = _build_history_x_labels(history)
    removed = [len(item["removed_ids"]) for item in history]
    fig = go.Figure()
    fig.add_trace(
        go.Bar(
            x=labels,
            y=removed,
            name="减少点位",
            marker_color="#8bb5d9",
            opacity=0.86,
        )
    )
    fig.update_layout(
        title=_centered_chart_title("减少点位变化"),
        height=300,
        margin={"l": 28, "r": 28, "t": 56, "b": 34},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.92)",
        hovermode="x unified",
    )
    fig.update_xaxes(type="category", tickangle=-22, categoryorder="array", categoryarray=labels)
    fig.update_yaxes(title_text="减少点位数", rangemode="tozero")
    return fig


def build_pollution_ratio_figure(history: List[Dict[str, Any]]):
    import plotly.graph_objects as go

    labels = _build_history_x_labels(history)
    ratios = [float(item["pollution_ratio"]) * 100 for item in history]
    fig = go.Figure()
    fig.add_trace(
        go.Scatter(
            x=labels,
            y=ratios,
            name="污染占比(%)",
            mode="lines+markers",
            line={"color": "#1f4e7a", "width": 2.8},
            marker={"size": 7, "color": "#ffffff", "line": {"width": 2, "color": "#1f4e7a"}},
        )
    )
    fig.update_layout(
        title=_centered_chart_title("污染占比趋势"),
        height=300,
        margin={"l": 28, "r": 28, "t": 56, "b": 34},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.92)",
        hovermode="x unified",
    )
    fig.update_xaxes(type="category", tickangle=-22, categoryorder="array", categoryarray=labels)
    fig.update_yaxes(title_text="污染占比(%)", rangemode="tozero")
    return fig


def build_centroid_drift_figure(
    history: List[Dict[str, Any]],
    show_all_dates: bool = False,
    smoothing_window: int = 1,
):
    import plotly.graph_objects as go

    centroid_points = build_centroid_window_points(history, smoothing_window)
    xs: List[float] = []
    ys: List[float] = []
    labels: List[str] = []
    sizes: List[float] = []
    hover_text: List[str] = []

    for point in centroid_points:
        row = point["row"]
        col = point["col"]
        xs.append(col)
        ys.append(ROWS - row + 1)
        labels.append(point["period_label"])
        sizes.append(min(28, 11 + point["polluted_count"] * 0.35))
        hover_text.append(
            f"时间窗口：{point['period_label']}<br>"
            f"代表日期：{point['representative_label']}<br>"
            f"代表污染点：ID {point['point_id']}（行{point['row']} / 列{point['col']}）<br>"
            f"平滑重心：行{point['target_row']:.1f} / 列{point['target_col']:.1f}<br>"
            f"污染点位：{point['polluted_count']}<br>连通块：{point['cluster_count']}"
        )

    fig = go.Figure()
    if xs:
        fig.add_trace(
            go.Scatter(
                x=xs,
                y=ys,
                mode="lines+markers",
                marker={
                    "size": sizes,
                    "color": list(range(len(xs))),
                    "colorscale": "Blues",
                    "line": {"width": 2, "color": "#ffffff"},
                    "opacity": 0.9,
                },
                line={"color": "#2e6ea9", "width": 3},
                hovertext=hover_text,
                hovertemplate="%{hovertext}<extra></extra>",
                name="污染重心",
                selected={"marker": {"size": 22, "color": "#1f5a92", "opacity": 1.0}},
                unselected={"marker": {"opacity": 0.86}},
            )
        )
        if show_all_dates:
            fig.add_trace(
                go.Scatter(
                    x=xs,
                    y=ys,
                    mode="text",
                    text=labels,
                    textposition="top center",
                    textfont={"size": 12, "color": "#6b7f96"},
                    hoverinfo="skip",
                    showlegend=False,
                )
            )
        else:
            start_label = f"起始 {labels[0]}"
            end_label = f"当前 {labels[-1]}"
            if len(xs) == 1:
                start_label = f"起始/当前 {labels[0]}"
                end_label = ""

            fig.add_trace(
                go.Scatter(
                    x=[xs[0]],
                    y=[ys[0]],
                    mode="markers+text",
                    text=[start_label],
                    textposition="top left",
                    marker={
                        "size": max(18, sizes[0] + 2),
                        "color": "#d3e6f8",
                        "line": {"width": 2.4, "color": "#2e6ea9"},
                    },
                    hovertext=[hover_text[0]],
                    hovertemplate="%{hovertext}<extra></extra>",
                    name="起始位置",
                )
            )
            if len(xs) > 1:
                fig.add_trace(
                    go.Scatter(
                        x=[xs[-1]],
                        y=[ys[-1]],
                        mode="markers+text",
                        text=[end_label],
                        textposition="top right",
                        marker={
                            "size": max(20, sizes[-1] + 3),
                            "color": "#2e6ea9",
                            "line": {"width": 2.6, "color": "#ffffff"},
                        },
                        textfont={"size": 13, "color": "#204d7c"},
                        hovertext=[hover_text[-1]],
                        hovertemplate="%{hovertext}<extra></extra>",
                        name="当前位置",
                    )
                )

    fig.update_layout(
        title=_centered_chart_title(f"污染重心漂移轨迹 | {max(1, int(smoothing_window))} 期窗口"),
        height=780,
        margin={"l": 24, "r": 20, "t": 68, "b": 36},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.92)",
        hovermode="closest",
        clickmode="event+select",
        legend={
            "orientation": "h",
            "yanchor": "bottom",
            "y": 1.02,
            "xanchor": "left",
            "x": 0.0,
        },
    )
    fig.update_xaxes(title_text="列", range=[0.5, COLS + 0.5], dtick=1, gridcolor="#e5eef8", zeroline=False)
    fig.update_yaxes(title_text="行", range=[0.5, ROWS + 0.5], dtick=1, gridcolor="#e5eef8", zeroline=False)
    return fig


def build_distribution_grid_figure(
    title: str,
    value_by_id: Dict[int, float],
    value_label: str,
    obstacle_ids: Set[int],
    pending_ids: Set[int],
    highlight_labels: Optional[Dict[int, str]] = None,
    colorscale: str = "YlOrRd",
):
    import plotly.graph_objects as go

    accessible_xs: List[int] = []
    accessible_ys: List[int] = []
    accessible_values: List[float] = []
    accessible_hover: List[str] = []
    obstacle_xs: List[int] = []
    obstacle_ys: List[int] = []
    pending_xs: List[int] = []
    pending_ys: List[int] = []
    max_value = max(value_by_id.values()) if value_by_id else 1.0
    max_value = max(max_value, 1.0)

    for point_id in range(1, TOTAL + 1):
        row, col = id_to_row_col(point_id)
        x = col
        y = ROWS - row + 1
        if point_id in obstacle_ids:
            obstacle_xs.append(x)
            obstacle_ys.append(y)
            continue
        if point_id in pending_ids:
            pending_xs.append(x)
            pending_ys.append(y)
            continue

        value = float(value_by_id.get(point_id, 0.0))
        accessible_xs.append(x)
        accessible_ys.append(y)
        accessible_values.append(value)
        accessible_hover.append(
            f"ID {point_id}<br>行{row} 列{col}<br>{value_label}: {value:.2f}"
        )

    fig = go.Figure()
    fig.add_trace(
        go.Scatter(
            x=accessible_xs,
            y=accessible_ys,
            mode="markers",
            marker={
                "symbol": "square",
                "size": 20,
                "color": accessible_values,
                "colorscale": colorscale,
                "cmin": 0,
                "cmax": max_value,
                "line": {"width": 0.35, "color": "rgba(255,255,255,0.65)"},
                "colorbar": {
                    "title": value_label,
                    "thickness": 14,
                    "len": 0.84,
                    "outlinewidth": 0,
                },
            },
            hovertext=accessible_hover,
            hovertemplate="%{hovertext}<extra></extra>",
            name=value_label,
        )
    )

    if pending_xs:
        fig.add_trace(
            go.Scatter(
                x=pending_xs,
                y=pending_ys,
                mode="markers",
                marker={"symbol": "square", "size": 20, "color": "#fff7cc", "line": {"width": 0.4, "color": "#d0b24d"}},
                hovertemplate="待测区域<extra></extra>",
                name="待测区域",
            )
        )

    if obstacle_xs:
        fig.add_trace(
            go.Scatter(
                x=obstacle_xs,
                y=obstacle_ys,
                mode="markers",
                marker={"symbol": "square", "size": 20, "color": "#7f8895", "line": {"width": 0.4, "color": "#ffffff"}},
                hovertemplate="障碍物<extra></extra>",
                name="障碍物",
            )
        )

    if highlight_labels:
        marker_xs: List[int] = []
        marker_ys: List[int] = []
        marker_texts: List[str] = []
        hover_texts: List[str] = []
        for point_id, label in highlight_labels.items():
            row, col = id_to_row_col(point_id)
            marker_xs.append(col)
            marker_ys.append(ROWS - row + 1)
            marker_texts.append(label)
            hover_texts.append(f"{label}<br>ID {point_id}<br>行{row} 列{col}")
        fig.add_trace(
            go.Scatter(
                x=marker_xs,
                y=marker_ys,
                mode="markers+text",
                text=marker_texts,
                textposition="middle center",
                marker={
                    "symbol": "star",
                    "size": 20,
                    "color": "#102b47",
                    "line": {"width": 2, "color": "#ffd166"},
                },
                textfont={"size": 10, "color": "#ffffff"},
                hovertext=hover_texts,
                hovertemplate="%{hovertext}<extra></extra>",
                name="重点标记",
            )
        )

    fig.update_layout(
        title=_centered_chart_title(title),
        height=740,
        margin={"l": 8, "r": 8, "t": 62, "b": 12},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.94)",
        legend={"orientation": "h", "yanchor": "top", "y": -0.11, "x": 0},
    )
    fig.update_xaxes(range=[0.5, COLS + 0.5], showgrid=False, visible=False)
    fig.update_yaxes(range=[0.5, ROWS + 0.5], showgrid=False, visible=False)
    return fig


def build_hotspot_bar_figure(history: List[Dict[str, Any]]):
    import plotly.graph_objects as go

    stats = build_point_history_stats(history)
    ranked = sorted(
        stats.items(),
        key=lambda item: (-item[1]["occurrences"], -item[1]["max_streak"], item[0]),
    )[:10]
    point_labels = [f"ID {point_id}" for point_id, _ in ranked]
    freq_values = [value["occurrences"] for _, value in ranked]
    streak_values = [value["max_streak"] for _, value in ranked]

    fig = go.Figure()
    fig.add_trace(
        go.Scatter(
            x=point_labels,
            y=freq_values,
            name="出现次数",
            mode="lines+markers",
            line={"color": "#2e6ea9", "width": 3},
            marker={"size": 8, "color": "#ffffff", "line": {"width": 2, "color": "#2e6ea9"}},
        )
    )
    fig.add_trace(
        go.Scatter(
            x=point_labels,
            y=streak_values,
            name="最长连续天数",
            mode="lines+markers",
            line={"color": "#e14b43", "width": 3},
            marker={"size": 8, "color": "#ffffff", "line": {"width": 2, "color": "#e14b43"}},
        )
    )
    fig.update_layout(
        title=_centered_chart_title("热点排名折线对比", y=0.975),
        height=380,
        margin={"l": 22, "r": 22, "t": 82, "b": 28},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.92)",
        xaxis={"tickangle": -24},
        yaxis={"title": "数值", "rangemode": "tozero"},
        legend={"orientation": "h", "yanchor": "bottom", "y": 1.09, "x": 0.5, "xanchor": "center"},
        hovermode="x unified",
    )
    return fig


def build_candidate_score_line_figure(
    candidates: List[Dict[str, Any]],
    title: str,
    value_label: str,
) -> Any:
    import plotly.graph_objects as go

    fig = go.Figure()
    if not candidates:
        fig.update_layout(
            title=_centered_chart_title(title),
            height=360,
            margin={"l": 22, "r": 22, "t": 50, "b": 28},
            paper_bgcolor="rgba(255,255,255,0)",
            plot_bgcolor="rgba(255,255,255,0.92)",
            annotations=[
                {
                    "text": "暂无可绘制的数据",
                    "xref": "paper",
                    "yref": "paper",
                    "x": 0.5,
                    "y": 0.5,
                    "showarrow": False,
                    "font": {"size": 15, "color": "#47698b"},
                }
            ],
        )
        return fig

    labels = [f"ID {int(item['point_id'])}" for item in candidates[:20]]
    values = [float(item["score"]) * 100 for item in candidates[:20]]
    fig.add_trace(
        go.Scatter(
            x=labels,
            y=values,
            mode="lines+markers",
            name=value_label,
            line={"color": "#2f6ea9", "width": 3},
            marker={"size": 8, "color": "#ffffff", "line": {"width": 2, "color": "#2f6ea9"}},
            hovertemplate="%{x}<br>" + value_label + "=%{y:.1f}<extra></extra>",
        )
    )
    fig.update_layout(
        title=_centered_chart_title(title),
        height=380,
        margin={"l": 22, "r": 22, "t": 50, "b": 28},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.92)",
        xaxis={"tickangle": -24},
        yaxis={"title": value_label, "rangemode": "tozero"},
        hovermode="x unified",
    )
    return fig


def render_analytics_section(
    report_options: List[Tuple[str, str]],
    report_store: Dict[str, Dict[str, Any]],
) -> None:
    import streamlit as st

    history = build_report_history(report_options, report_store)
    if not history:
        render_note_bar("暂无可用于分析的历史数据。", soft=True)
        return

    latest = history[-1]
    st.subheader("五、趋势预测分析")
    analysis_config = default_analysis_config()
    forecast_count = forecast_next_pollution_count(history, analysis_config)
    frequency_map = build_frequency_map(history)
    point_stats = build_point_history_stats(history)
    risk_candidates = build_risk_candidates(history, analysis_config)
    source_candidates = build_source_candidates(history, analysis_config)

    stable_hotspots = sum(1 for value in point_stats.values() if value["max_streak"] >= 2)
    avg_count = round(sum(item["polluted_count"] for item in history) / len(history))
    latest_change = len(latest["added_ids"]) - len(latest["removed_ids"])

    render_analytics_hero(
        latest,
        forecast_count,
        stable_hotspots,
        risk_candidates[0] if risk_candidates else None,
        source_candidates[0] if source_candidates else None,
    )
    render_analytics_section_map()

    k1, k2, k3, k4 = st.columns(4)
    with k1:
        render_metric_card("历史报告数", len(history))
    with k2:
        render_metric_card("平均污染点位", avg_count, alert=avg_count > 0)
    with k3:
        render_metric_card("预测下一期", forecast_count, alert=forecast_count >= latest["polluted_count"])
    with k4:
        render_metric_card("持续热点数", stable_hotspots, alert=stable_hotspots > 0)

    st.markdown('<div class="section-gap-sm"></div>', unsafe_allow_html=True)
    analysis_config = render_analysis_config_panel()
    forecast_count = forecast_next_pollution_count(history, analysis_config)
    risk_candidates = build_risk_candidates(history, analysis_config)
    source_candidates = build_source_candidates(history, analysis_config)
    tabs = st.tabs(["趋势总览", "热点画像", "趋势预测", "污染溯源"])

    with tabs[0]:
        render_analytics_subsection(
            "趋势总览",
            "这一层先回答污染总量怎么变、扩张还是收缩、空间重心是否在移动。",
            "5.1",
        )
        render_note_bar(
            "趋势总览按日期维度拆分为多张图，分别展示污染点位、新增点位、减少点位、污染占比和空间重心，避免多指标叠在同一坐标轴。"
            + f' 当前最近一期连通块数量为 {latest["cluster_count"]}，说明污染分布'
            + ("更集中。" if latest["cluster_count"] <= 2 else "较为分散。")
        )
        render_analysis_methodology(
            "5.1 模块说明与指标解释",
            [
                "趋势总览模块用于从时序角度刻画污染规模与空间结构的演化过程，其核心任务不是给出单一结论，而是同时回答污染总量是否变化、变化方向是否稳定，以及污染重心是否发生显著迁移。",
                "在解释图形结果时，应结合污染点位总数、新增与减少点位、污染占比以及连通块数量进行联合判断。若总量上升且重心持续偏移，通常意味着污染范围存在扩张倾向；若总量下降但结构仍高度集中，则说明污染可能转入局部残留阶段。",
                "本模块中的“日期地图快照”用于提供空间参照，“重心漂移轨迹”用于观察空间迁移趋势，下方多张时序图则分别承担总量、增量和比例的独立表达，从而避免单图多轴带来的解释歧义。",
            ],
            formulas=[
                r"\Delta N_t = N_t - N_{t-1}",
                r"P_t = \frac{N_t}{A_t}",
                r"C_t = \left(\frac{1}{N_t}\sum_{i=1}^{N_t} r_i,\ \frac{1}{N_t}\sum_{i=1}^{N_t} c_i\right)",
            ],
            formula_explanations=[
                [
                    "符号说明：N_t 表示第 t 期的污染点位总数，N_{t-1} 表示上一期的污染点位总数，ΔN_t 表示相邻两期之间的净变化量。",
                    "解释说明：如果 ΔN_t > 0，说明这一期比上一期污染点位更多，整体呈扩张趋势；如果 ΔN_t < 0，则表示污染规模在收缩。",
                ],
                [
                    "符号说明：P_t 表示第 t 期污染占比，N_t 为该期污染点位数，A_t 为该期可参与判断的总点位数，即总网格数扣除障碍物与待测区域后的可达点位数。",
                    "解释说明：这个指标反映的是“可用区域里有多少比例被判为污染”，比只看总数更适合做跨日期比较，因为它剔除了障碍和待测区变化带来的干扰。",
                ],
                [
                    "符号说明：C_t 为第 t 期污染重心坐标；r_i 表示第 i 个污染点位所在行号，c_i 表示所在列号；求和后再除以 N_t，得到所有污染点位在行、列两个方向上的平均位置。",
                    "解释说明：重心不是某一个真实点位，而是所有污染点位的“平均中心位置”。如果这个中心持续向某个方向移动，通常说明污染重点区域正在发生空间迁移。",
                ],
            ],
        )
        map_labels = _build_history_x_labels(history)
        map_index = st.select_slider(
            "日期地图切换",
            options=list(range(len(history))),
            value=len(history) - 1,
            format_func=lambda idx: map_labels[int(idx)],
            key="overview_map_date_slider",
        )
        st.markdown('<div class="analytics-tab-block-gap"></div>', unsafe_allow_html=True)
        st.plotly_chart(
            build_analytics_date_map_figure(history[int(map_index)]),
            use_container_width=True,
            config={
                "staticPlot": True,
                "responsive": True,
                "displaylogo": False,
                "displayModeBar": False,
            },
        )
        centroid_window_cap = 10 if len(history) <= 10 else 20
        centroid_window_max = max(1, min(len(history), centroid_window_cap))
        recommended_centroid_window = recommend_centroid_smoothing_window(len(history), centroid_window_max)
        render_note_bar(
            "下方重心追踪图支持窗口平滑。你可以直接拖动窗口滑块控制轨迹密度，并在右侧切换是否实时显示全部日期标注。",
            soft=True,
        )
        render_note_bar(
            f"当前共收集 {len(history)} 期历史，系统建议使用 {recommended_centroid_window} 期窗口。"
            " 平滑后的代表性重心会吸附到真实出现过污染的点位，避免落在无污染坐标上。",
            soft=True,
        )
        control_col, toggle_col = st.columns([1.8, 1.0], gap="medium")
        with control_col:
            if centroid_window_max <= 1:
                centroid_smoothing_window = 1
                st.caption("重心平滑窗口（按期）：当前仅 1 期历史数据，窗口固定为 1。")
            else:
                centroid_smoothing_window = st.slider(
                    "重心平滑窗口（按期）",
                    1,
                    centroid_window_max,
                    recommended_centroid_window,
                    key="overview_centroid_smoothing_window",
                    help="窗口为 1 表示不平滑；窗口越大，系统会把连续多期报告压缩为一个代表性重心，并吸附到窗口内真实存在污染的点位。",
                )
        with toggle_col:
            show_centroid_dates = st.toggle(
                "实时显示重心日期标注",
                value=False,
                key="overview_centroid_show_all_dates",
                help="关闭时仅标注起始位置和当前位置；开启后会在轨迹上的每个日期点实时显示日期。",
            )
        smoothed_centroid_points = build_centroid_window_points(history, centroid_smoothing_window)
        render_note_bar(
            f"当前窗口为 {centroid_smoothing_window} 期，轨迹将展示 {len(smoothed_centroid_points)} 个代表性重心。"
            " 鼠标悬浮任一点可查看对应窗口范围、代表日期、真实污染点坐标和平滑后的中心位置。",
            soft=True,
        )
        st.markdown('<div class="analytics-tab-block-gap"></div>', unsafe_allow_html=True)
        st.plotly_chart(
            build_centroid_drift_figure(
                history,
                show_all_dates=show_centroid_dates,
                smoothing_window=centroid_smoothing_window,
            ),
            use_container_width=True,
            config={
                "responsive": True,
                "scrollZoom": False,
                "displaylogo": False,
                "doubleClick": False,
            },
        )
        render_analytics_mini_divider()
        st.plotly_chart(build_pollution_count_figure(history, forecast_count), use_container_width=True)
        st.plotly_chart(build_added_points_figure(history), use_container_width=True)
        st.plotly_chart(build_removed_points_figure(history), use_container_width=True)
        st.plotly_chart(build_pollution_ratio_figure(history), use_container_width=True)

        overview_rows = [
            {
                "日期": item["date"],
                "污染点位数": item["polluted_count"],
                "污染占比": f"{item['pollution_ratio'] * 100:.1f}%",
                "新增点位数": len(item["added_ids"]),
                "减少点位数": len(item["removed_ids"]),
                "连通块数": item["cluster_count"],
                "重心": (
                    f"行{item['centroid'][0]:.1f} / 列{item['centroid'][1]:.1f}"
                    if item["centroid"]
                    else "-"
                ),
            }
            for item in history
        ]
        render_export_button("导出趋势总览 CSV", overview_rows, "trend_overview.csv", "export_trend_overview")
        st.dataframe(overview_rows, use_container_width=True, hide_index=True)
        render_note_bar(
            f"当前最新一期相较前一期净变化 {latest_change:+d} 个点位；预测值采用近 4 期加权平均并叠加近期漂移。",
            soft=True,
        )

    with tabs[1]:
        st.markdown('<div class="analytics-tab-top-gap"></div>', unsafe_allow_html=True)
        render_analytics_subsection(
            "热点画像",
            "这一层把历史热点和持续性拆开看，帮助判断哪些点是长期问题，哪些只是偶发出现。",
            "5.2",
        )
        render_note_bar("热点画像统计的是“哪里反复出现、哪里持续最久”。如果想更强调持续污染而不是出现频次，可以在上方参数里提高溯源和风险中的连续性权重。")
        render_analysis_methodology(
            "5.2 模块说明与统计逻辑",
            [
                "热点画像模块用于识别具有重复暴露特征的重点点位，其分析重点在于区分“高频出现”与“长期持续”这两类不同风险。前者反映历史复现性，后者反映污染稳定性，二者在现场处置中对应不同的复核优先级。",
                "热力分布图用于描述历史上各点位的累积活跃程度，右侧排行表则将频次、最长连续天数和最近连续天数进行结构化展开。若某点位同时具有高频与高持续特征，通常可视为稳定热点或长期异常区域。",
                "该模块适合用于形成重点观察名单，为后续趋势预测和污染溯源提供先验分布依据；因此其结果更偏统计表征，而非直接给出处置结论。",
            ],
            formulas=[
                r"f_i = \frac{\text{occurrences}_i}{T}",
                r"H_i = \alpha f_i + \beta s_i + \gamma c_i",
            ],
            formula_explanations=[
                [
                    "符号说明：f_i 表示点位 i 的历史出现频率，occurrences_i 表示该点位在全部历史期次中被识别为污染的次数，T 表示总期数。",
                    "解释说明：f_i 越大，说明这个点位在历史上越常出现，属于“高频热点”；如果某点只出现过一次，即使那次很严重，频率值也不会太高。",
                ],
                [
                    "符号说明：H_i 表示点位 i 的热点综合得分；α、β、γ 为权重系数；f_i 表示历史频率，s_i 表示持续性指标（如最长连续期数），c_i 表示当前或最近阶段的活跃程度。",
                    "解释说明：这个公式是在把“出现得多”“持续得久”“最近是否还活跃”三类信息合并成一个总分。最终分数越高，越值得被视作长期关注热点。",
                ],
            ],
        )
        ranked_points = sorted(
            point_stats.items(),
            key=lambda item: (-item[1]["occurrences"], -item[1]["max_streak"], item[0]),
        )[:12]
        hotspot_rows = [
            {
                "点位ID": point_id,
                "出现次数": value["occurrences"],
                "出现频率": f"{value['frequency_ratio'] * 100:.1f}%",
                "最长连续天数": value["max_streak"],
                "最近连续天数": value["current_streak"],
            }
            for point_id, value in ranked_points
        ]
        st.markdown('<div class="analytics-tab-block-gap"></div>', unsafe_allow_html=True)
        hotspot_col, ranking_col = st.columns([1.34, 1.26], gap="large")
        with hotspot_col:
            st.plotly_chart(
                build_distribution_grid_figure(
                    "历史热点热力分布",
                    {point_id: float(count) for point_id, count in frequency_map.items()},
                    "历史出现次数",
                    latest["obstacle_ids"],
                    latest["pending_ids"],
                    colorscale=[
                        [0.00, "#f6f7f2"],
                        [0.24, "#eae4cf"],
                        [0.48, "#d8c89a"],
                        [0.72, "#b48e54"],
                        [1.00, "#74502a"],
                    ],
                ),
                use_container_width=True,
            )
        with ranking_col:
            render_export_button("导出热点排行 CSV", hotspot_rows, "hotspot_ranking.csv", "export_hotspot_rows")
            st.dataframe(hotspot_rows, use_container_width=True, hide_index=True, height=680)
        render_analytics_mini_divider()
        st.plotly_chart(build_hotspot_bar_figure(history), use_container_width=True)

    with tabs[2]:
        st.markdown('<div class="analytics-tab-top-gap"></div>', unsafe_allow_html=True)
        render_analytics_subsection(
            "趋势预测",
            "这一层基于截至当前期的历史快照推演下一期风险，适合用来做提前复核和资源安排。",
            "5.3",
        )
        render_note_bar("趋势预测关注的是“下一期哪里更危险”。可通过上方参数面板调高历史频率、邻域支撑或当前污染加分，观察预测热区如何变化。")
        render_analysis_methodology(
            "5.3 模块说明与预测思想",
            [
                "趋势预测模块面向下一期研判，其目标是根据截至当前时点的历史记录，对下一期潜在高风险点位进行排序，而不是直接生成确定性标签。因而本模块输出的是风险热度与候选序列，适用于提前复核、资源布置和重点巡检。",
                "预测逻辑综合考虑历史频率、近几期活跃度、当前连续性、邻域支撑以及区域扩散特征。若某点位自身活跃度高、周边已存在污染支撑且位于扩散边界，则其下一期风险通常更高。",
                "回放模式允许用户在不同历史截面重复计算预测结果，从而比较模型在各时点下的判断差异。这一设计有助于验证预测稳定性，并支持参数校准与方法复盘。",
            ],
            formulas=[
                r"R_i = w_f F_i + w_r R_i^{(recent)} + w_s S_i + w_n N_i + b_i",
                r"\hat{N}_{t+1} = \operatorname{Smooth}(N_{t-k:t}) + \lambda \cdot \operatorname{Drift}_t",
            ],
            formula_explanations=[
                [
                    "符号说明：R_i 表示点位 i 的综合风险分；w_f、w_r、w_s、w_n 分别表示历史频率、近期活跃度、连续性和邻域支撑的权重；F_i 表示历史频率，R_i^(recent) 表示近几期活跃程度，S_i 表示连续污染强度，N_i 表示周边污染支撑，b_i 表示附加修正项。",
                    "解释说明：这个分数越高，表示系统越倾向于认为该点位在下一期继续出现或扩散的风险更大。它不是单看某一项，而是把“历史、近期、连续性、周边环境”一起考虑。",
                ],
                [
                    "符号说明：N̂_(t+1) 表示对下一期污染点位总数的预测值；Smooth(N_(t-k:t)) 表示对最近 k 期总量进行平滑后的结果；Drift_t 表示最近阶段的变化趋势；λ 为趋势放大量级。",
                    "解释说明：先用最近几期的平均走势估计一个基础值，再根据最近是上升还是下降做适度修正，因此该值本质上是“平滑基线 + 趋势补偿”的组合结果。",
                ],
            ],
        )
        prediction_playback_index = render_playback_controls(
            history,
            "prediction_playback_slider",
            "按日期回放预测推演",
        )
        prediction_snapshot = build_analytics_snapshot(history, prediction_playback_index, analysis_config)
        prediction_latest = prediction_snapshot["latest"]
        prediction_risk_candidates = prediction_snapshot["risk_candidates"]
        prediction_forecast_count = prediction_snapshot["forecast_count"]
        prediction_finished = bool(st.session_state.get("prediction_playback_slider_finished"))
        render_note_bar(
            f"当前回放到 {prediction_latest['date']}，系统只使用前 {prediction_playback_index + 1} 期历史做下一期预测。",
            soft=True,
        )
        top_risk_rows = [
            {
                "点位ID": item["point_id"],
                "位置": f"行{item['row']} 列{item['col']}",
                "风险分": f"{item['score'] * 100:.1f}",
                "状态": item["current_state"],
                "历史频率": f"{item['frequency_ratio'] * 100:.1f}%",
                "近三期出现": f"{item['recent_ratio'] * 100:.1f}%",
                "邻域支撑": f"{item['neighbor_support'] * 100:.1f}%",
            }
            for item in prediction_risk_candidates[:15]
        ]
        if prediction_risk_candidates:
            top_risk = prediction_risk_candidates[0]
            render_note_bar(
                f"结论先行：按 {prediction_latest['date']} 当时的历史数据看，下一期重点关注 ID {top_risk['point_id']}，"
                f"位于行{top_risk['row']} 列{top_risk['col']}，综合风险分 {top_risk['score'] * 100:.1f}。",
            )
        if prediction_finished and prediction_risk_candidates:
            top_risk = prediction_risk_candidates[0]
            render_playback_completion_card(
                "预测回放已到最终态",
                f"最终停留在 {prediction_latest['date']}。系统给出的下一期首要关注点位是 "
                f"ID {top_risk['point_id']}（行{top_risk['row']} 列{top_risk['col']}），"
                f"风险分 {top_risk['score'] * 100:.1f}。",
            )
        st.markdown('<div class="analytics-tab-block-gap"></div>', unsafe_allow_html=True)
        prediction_col, table_col = st.columns([1.34, 1.26], gap="large")
        with prediction_col:
            risk_value_map = {item["point_id"]: item["score"] * 100 for item in prediction_risk_candidates[:160]}
            st.plotly_chart(
                build_distribution_grid_figure(
                    f"下一期污染风险热力图 | 回放到 {prediction_latest['date']}",
                    risk_value_map,
                    "风险分",
                    prediction_latest["obstacle_ids"],
                    prediction_latest["pending_ids"],
                    colorscale="Sunset",
                ),
                use_container_width=True,
            )
        with table_col:
            top_risk_rows = [
                {
                    "点位ID": item["point_id"],
                    "位置": f"行{item['row']} 列{item['col']}",
                    "风险分": f"{item['score'] * 100:.1f}",
                    "状态": item["current_state"],
                    "历史频率": f"{item['frequency_ratio'] * 100:.1f}%",
                    "近三期出现": f"{item['recent_ratio'] * 100:.1f}%",
                    "邻域支撑": f"{item['neighbor_support'] * 100:.1f}%",
                }
                for item in prediction_risk_candidates[:15]
            ]
            render_export_button(
                "导出风险排行 CSV",
                top_risk_rows,
                "prediction_risk_ranking.csv",
                "export_prediction_risk_rows",
            )
            st.dataframe(top_risk_rows, use_container_width=True, hide_index=True, height=680)
        render_analytics_mini_divider()
        st.plotly_chart(
            build_candidate_score_line_figure(
                prediction_risk_candidates,
                f"风险点位评分折线 | 回放到 {prediction_latest['date']}",
                "风险分",
            ),
            use_container_width=True,
        )
        p1, p2, p3 = st.columns(3)
        with p1:
            render_metric_card("回放期数", prediction_playback_index + 1)
        with p2:
            render_metric_card("当期污染点位", prediction_latest["polluted_count"], alert=True)
        with p3:
            render_metric_card(
                "当时预测下一期",
                prediction_forecast_count,
                alert=prediction_forecast_count >= prediction_latest["polluted_count"],
            )
        render_note_bar("风险分结合历史出现频率、近三期活跃度、当前连续天数和周边已污染邻域强度。", soft=True)

    with tabs[3]:
        st.markdown('<div class="analytics-tab-top-gap"></div>', unsafe_allow_html=True)
        render_analytics_subsection(
            "污染溯源",
            "这一层强调候选排序而不是唯一答案，用来辅助现场决定先排查哪里。",
            "5.4",
        )
        render_note_bar("污染溯源给出的不是唯一结论，而是“更值得优先排查的候选点”。如果现场更看重最早出现还是持续时间更长，可以直接调上方溯源权重。")
        render_analysis_methodology(
            "5.4 模块说明与源点判别依据",
            [
                "污染溯源模块用于生成疑似源点候选序列，其本质是对“最早出现、持续时间较长、后续带动周边扩散且接近热点中心”的点位进行综合评分。该模块并不将单一点位直接认定为唯一源头，而是形成可供人工复核的优先排查列表。",
                "在分析解释上，应同时关注首次出现期次、最长连续期数、扩散带动数量以及与整体热点中心的空间关系。若某点位在较早期即出现，且随后周边点位陆续活跃，则其更可能具备源头特征。",
                "回放溯源判断的意义在于检验候选源点是否具有时间一致性。若某点位在多个回放截面中持续排名靠前，则说明其源头解释具有更高稳健性和可复核价值。",
            ],
            formulas=[
                r"S_i = w_a A_i + w_p P_i + w_d D_i + w_c C_i",
                r"A_i \propto \frac{1}{1 + t_i^{first}}",
            ],
            formula_explanations=[
                [
                    "符号说明：S_i 表示点位 i 的源点评分；w_a、w_p、w_d、w_c 分别表示首次出现优势、持续性、扩散能力和中心关联性的权重；A_i 表示首次出现优势，P_i 表示持续污染强度，D_i 表示后续扩散带动能力，C_i 表示与热点中心的空间关系。",
                    "解释说明：某个点位如果出现得早、持续得久、后面又带动了周边点位活跃，同时还位于热点核心区域附近，那么它的综合源点评分就会更高。",
                ],
                [
                    "符号说明：t_i^first 表示点位 i 首次出现的期次编号；A_i 与 1 / (1 + t_i^first) 成正比，意味着首次出现越早，对应的 A_i 越大。",
                    "解释说明：越早出现的点位，越有可能是后续污染演化的起始位置，因此系统会给更高的“首发优势”分值；越晚才出现的点位，通常更像是被扩散影响到的结果点。",
                ],
            ],
        )
        source_playback_index = render_playback_controls(
            history,
            "source_playback_slider",
            "按日期回放溯源判断",
        )
        source_snapshot = build_analytics_snapshot(history, source_playback_index, analysis_config)
        source_latest = source_snapshot["latest"]
        playback_source_candidates = source_snapshot["source_candidates"]
        source_finished = bool(st.session_state.get("source_playback_slider_finished"))
        source_rows = [
            {
                "排名": f"S{rank + 1}",
                "点位ID": item["point_id"],
                "位置": f"行{item['row']} 列{item['col']}",
                "源点评分": f"{item['score'] * 100:.1f}",
                "首次出现": f"第 {item['first_seen_index'] + 1} 期",
                "最长连续": item["max_streak"],
                "后续扩散": item["spread_targets"],
            }
            for rank, item in enumerate(playback_source_candidates[:8])
        ]
        render_note_bar(
            f"当前回放到 {source_latest['date']}，系统会用截至该期的历史记录重算疑似源点。",
            soft=True,
        )
        if playback_source_candidates:
            top_source = playback_source_candidates[0]
            render_note_bar(
                f"结论先行：按 {source_latest['date']} 当时的历史数据看，最值得优先排查的疑似源点是 "
                f"ID {top_source['point_id']}（行{top_source['row']} 列{top_source['col']}），评分 {top_source['score'] * 100:.1f}。",
            )
        if source_finished and playback_source_candidates:
            top_source = playback_source_candidates[0]
            render_playback_completion_card(
                "溯源回放已到最终态",
                f"最终停留在 {source_latest['date']}。综合首次出现、持续性与扩散链路后，"
                f"当前最值得优先排查的疑似源点是 ID {top_source['point_id']} "
                f"（行{top_source['row']} 列{top_source['col']}），评分 {top_source['score'] * 100:.1f}。",
            )
        st.markdown('<div class="analytics-tab-block-gap"></div>', unsafe_allow_html=True)
        source_col, explanation_col = st.columns([1.34, 1.26], gap="large")
        with source_col:
            source_value_map = {item["point_id"]: item["score"] * 100 for item in playback_source_candidates[:160]}
            highlight_labels = {
                item["point_id"]: f"S{rank + 1}"
                for rank, item in enumerate(playback_source_candidates[:5])
            }
            st.plotly_chart(
                build_distribution_grid_figure(
                    f"疑似污染源分布图 | 回放到 {source_latest['date']}",
                    source_value_map,
                    "源点评分",
                    source_latest["obstacle_ids"],
                    source_latest["pending_ids"],
                    highlight_labels=highlight_labels,
                    colorscale="Tealrose",
                ),
                use_container_width=True,
            )
        with explanation_col:
            render_export_button(
                "导出溯源排行 CSV",
                source_rows,
                "source_ranking.csv",
                "export_source_rows",
            )
            st.dataframe(source_rows, use_container_width=True, hide_index=True, height=680)
        render_analytics_mini_divider()
        st.plotly_chart(
            build_candidate_score_line_figure(
                playback_source_candidates,
                f"疑似源点评分折线 | 回放到 {source_latest['date']}",
                "源点评分",
            ),
            use_container_width=True,
        )
        s1, s2, s3 = st.columns(3)
        with s1:
            render_metric_card("回放期数", source_playback_index + 1)
        with s2:
            render_metric_card("当期污染点位", source_latest["polluted_count"], alert=True)
        with s3:
            render_metric_card("候选源点数", len(playback_source_candidates), alert=len(playback_source_candidates) > 0)
        if playback_source_candidates:
            top_source = playback_source_candidates[0]
            render_note_bar(top_source["explanation"], soft=True)
        render_note_bar("源点评分不是确定结论，而是根据首次出现时间、持续性、周边后续扩散和整体热点中心关系给出的候选排序。", soft=True)


def main():
    import streamlit as st

    st.set_page_config(page_title="辐射超标地面点位与手套孔研判平台", layout="wide")
    inject_styles()

    if "report_store" not in st.session_state:
        st.session_state.report_store = {}
    if "upload_widget_version" not in st.session_state:
        st.session_state.upload_widget_version = 0
    report_store: Dict[str, Dict[str, Any]] = st.session_state.report_store

    render_page_hero(
        [
            (
                "一",
                "报告接入与文本校核",
                "支持检测报告接入、日期识别与文本校核，",
                "用于建立多日期报告索引并同步识别污染点位与手套孔信息。",
            ),
            (
                "二",
                "区域定义与修正",
                "提供网格点选、框选与区域状态修正能力，",
                "用于维护污染、清洁、障碍和待测区域的空间定义结果。",
            ),
            (
                "三",
                "日期结果展示",
                "按当前激活日期展示统计概览、结果明细与地图分布，",
                "用于核查污染点位、手套孔信息及对应空间位置。",
            ),
            (
                "四",
                "趋势分析与研判",
                "提供趋势总览、热点画像、风险预测和污染溯源分析，",
                "用于支撑多日期结果对比、风险判断与后续排查决策。",
            ),
        ]
    )

    render_stage_panel_start(
        "数据接入总览",
        "一、报告接入与文本校核",
        "统一接入多份检测报告，自动抽取日期与污染信息，并支持在页面内直接修正报告文本。",
        chips=[
            ("接入方式", "多报告并行", "上传后会自动建立日期索引"),
            ("支持格式", "TXT / DOCX / PDF", "可混合上传并保留原始文本"),
            ("校核模式", "自动识别 + 人工修正", "修改文本后会实时同步污染点位"),
        ],
    )
    render_section_path_map(
        [
            ("1.1", "文件接入", "上传多份报告、建立日期索引，并切换查看当前激活报告。"),
            ("1.2", "文本校核", "直接修正报告文本，识别结果会同步更新。"),
        ]
    )
    render_section_subsection(
        "文件接入",
        "这一部分负责建立报告集合、识别日期并切换当前激活报告，所有接入相关操作都集中在这里。",
        "1.1",
    )
    render_stage_subpanel_start("文件接入", "1.1")
    uploader_key = f"report_uploader_{st.session_state.upload_widget_version}"
    uploaded_files = st.file_uploader(
        "上传检测报告（txt/docx/pdf，可多选）",
        type=["txt", "docx", "pdf"],
        accept_multiple_files=True,
        key=uploader_key,
    )
    if st.button(
        "一键清空上传文件",
        key="clear_uploaded_files",
        disabled=not uploaded_files,
        type="secondary",
    ):
        st.session_state.upload_widget_version += 1
        st.rerun()
    render_stage_subpanel_end()

    sample = """普查机器人地面及手套孔检测报告
辐射超标地面点位ID
131
718
748
辐射超标手套孔名称
手套孔1#
手套孔2#
"""

    current_keys: List[str] = []
    if uploaded_files:
        if "__sample__" in report_store:
            del report_store["__sample__"]

        for uploaded in uploaded_files:
            file_bytes = uploaded.getvalue()
            file_key = build_file_signature(uploaded.name, file_bytes)
            current_keys.append(file_key)

            if file_key not in report_store:
                initial_text = ""
                load_error = None
                try:
                    initial_text = extract_text_from_file(uploaded.name, file_bytes)
                except Exception as exc:
                    load_error = str(exc)

                task_end_date = extract_task_end_date(initial_text)
                report_date = (
                    task_end_date
                    or extract_date_from_text(initial_text)
                    or extract_date_from_text(uploaded.name)
                )
                report_store[file_key] = {
                    "file_name": uploaded.name,
                    "date": report_date,
                    "task_end_date": task_end_date,
                    "report_text": initial_text,
                    "load_error": load_error,
                    "obstacle_ids": default_obstacle_ids(),
                    "pending_ids": set(),
                    "polluted_ids": set(parse_report_text(initial_text).floor_ids),
                    "pollution_source_signature": build_text_signature(initial_text),
                    "selected_ids": set(),
                }
            else:
                report_store[file_key]["file_name"] = uploaded.name

        for stale_key in list(report_store.keys()):
            if stale_key not in current_keys:
                del report_store[stale_key]

        missing_task_end_files = [
            str(report_store[key].get("file_name", "-"))
            for key in current_keys
            if not report_store[key].get("task_end_date")
        ]
        if missing_task_end_files:
            st.warning(
                "以下文件未识别到“任务结束时间”，请确认报告内容或手工核对日期："
                + "、".join(missing_task_end_files)
            )
    else:
        report_store.clear()
        report_store["__sample__"] = {
            "file_name": "sample.txt",
            "date": "示例",
            "task_end_date": None,
            "report_text": sample,
            "load_error": None,
            "obstacle_ids": default_obstacle_ids(),
            "pending_ids": set(),
            "polluted_ids": set(parse_report_text(sample).floor_ids),
            "pollution_source_signature": build_text_signature(sample),
            "selected_ids": set(),
        }
        current_keys = ["__sample__"]

    for key in current_keys:
        report = report_store[key]
        if "obstacle_ids" not in report:
            report["obstacle_ids"] = default_obstacle_ids()
        if "pending_ids" not in report:
            report["pending_ids"] = set()
        if "selected_ids" not in report:
            report["selected_ids"] = set()
        if "polluted_ids" not in report:
            report["polluted_ids"] = set(parse_report_text(str(report.get("report_text", ""))).floor_ids)
        if "pollution_source_signature" not in report:
            report["pollution_source_signature"] = build_text_signature(str(report.get("report_text", "")))
        if "report_text" not in report:
            report["report_text"] = ""
        if "date" not in report:
            report["date"] = None
        if "task_end_date" not in report:
            report["task_end_date"] = extract_task_end_date(str(report.get("report_text", "")))
        if "load_error" not in report:
            report["load_error"] = None

    report_options = build_report_switch_options(current_keys, report_store)
    option_keys = [key for key, _ in report_options]
    if "active_report_key" not in st.session_state or st.session_state.active_report_key not in option_keys:
        st.session_state.active_report_key = option_keys[0]

    render_stage_meta_strip(
        [
            ("当前报告数", str(len(current_keys))),
            ("数据状态", "已上传报告" if uploaded_files else "示例数据"),
            ("当前激活", str(report_store[st.session_state.active_report_key].get("file_name", "-"))),
            (
                "识别日期",
                str(report_store[st.session_state.active_report_key].get("date") or "未识别日期"),
            ),
        ]
    )
    render_switcher_panel(
        "按日期切换报告文本",
        "top_report_switch",
        report_options,
        st.session_state.active_report_key,
    )
    active_report_key = st.session_state.active_report_key
    active_report = report_store[active_report_key]
    if active_report.get("load_error"):
        st.error(f"文件解析失败：{active_report.get('load_error')}")
    render_stage_subpanel_end()

    render_section_subsection(
        "文本校核",
        "这一部分用于核对原始报告文本和识别结果是否一致，必要时可以直接在页面中修正。",
        "1.2",
    )
    render_stage_subpanel_start("文本校核", "1.2")
    report_text_widget_key = f"report_text_{active_report_key}"
    report_text_notice_key = f"report_text_applied_notice_{active_report_key}"
    if report_text_widget_key not in st.session_state:
        st.session_state[report_text_widget_key] = active_report["report_text"]
    if st.session_state.pop(report_text_notice_key, False):
        st.success("报告文本修改已应用，识别结果与统计结果已同步更新。")
    draft_report_text = st.text_area(
        "检测报告文本（可手工修改）",
        key=report_text_widget_key,
        height=220,
    )
    if draft_report_text != str(active_report.get("report_text", "")):
        render_note_bar("当前文本存在尚未应用的修改。确认无误后，请点击下方“应用修改”按钮同步更新识别结果。", soft=True)
    if st.button("应用修改", key=f"apply_report_text_{active_report_key}", type="primary"):
        active_report["report_text"] = draft_report_text
        task_end_date = extract_task_end_date(draft_report_text)
        active_report["task_end_date"] = task_end_date
        active_report["date"] = (
            task_end_date
            or extract_date_from_text(draft_report_text)
            or extract_date_from_text(str(active_report.get("file_name", "")))
        )
        st.session_state[report_text_notice_key] = True
        st.rerun()
    parsed = sync_report_pollution_state(active_report)
    if parsed.floor_count_mismatch and parsed.declared_floor_count is not None:
        st.warning(
            f"报告声明的污染地面点位数为 {parsed.declared_floor_count}，"
            f"但当前识别/修正结果为 {len(get_report_polluted_ids(active_report))}。"
            "建议核对原始报告文本或人工修正结果。"
        )
    render_stage_subpanel_end()
    render_stage_panel_end()
    render_major_section_divider()

    render_stage_panel_start(
        "区域修正流程",
        "二、区域定义与修正",
        "在地图上统一完成区域基线设置、点选框选修正和污染状态写入，不再把障碍物初始化拆成单独一段。",
        chips=[
            (
                "当前地图日期",
                str(active_report.get("date") or "未识别日期"),
                "地图、结果和后续汇总都围绕当前激活日期联动切换",
            ),
            (
                "当前障碍点位",
                str(len(active_report["obstacle_ids"])),
                "固定阻挡区，不参与可达污染判断",
            ),
            (
                "待测区域",
                str(len(active_report["pending_ids"])),
                "等待复核，不计入污染或无污染",
            ),
        ],
    )
    render_section_path_map(
        [
            ("2.1", "区域基线设置", "先确定默认障碍、手工障碍或掩膜导入的基础区域。"),
            ("2.2", "地图修正", "再在地图上框选、点选并观察颜色预览后的区域变化。"),
            ("2.3", "状态写入", "最后把选区确认写入污染、清洁、障碍或待测状态。"),
        ]
    )
    render_stage_meta_strip(
        [
            ("区域来源", "默认模板 / 手工输入 / 掩膜导入"),
            (
                "当前可达区域",
                str(max(TOTAL - len(active_report["obstacle_ids"]) - len(active_report["pending_ids"]), 0)),
            ),
        ]
    )
    render_section_subsection(
        "区域基线设置",
        "这一层决定地图上的基础可达区域，适合在开始修正之前先把底图状态准备好。",
        "2.1",
    )
    render_stage_subpanel_start("区域基线设置", "2.1")
    init_mode = st.radio(
        "障碍物来源",
        ["使用默认障碍物", "手工输入障碍物ID", "上传掩膜图（黑障碍/白可达）"],
        horizontal=True,
        key=f"init_mode_{active_report_key}",
    )

    if init_mode == "手工输入障碍物ID":
        raw = st.text_area(
            "输入障碍物点位 ID（支持空格/逗号/换行分隔）",
            value="",
            key=f"manual_obstacle_{active_report_key}",
        )
        if st.button("应用手工障碍物", key=f"apply_manual_{active_report_key}", type="primary"):
            active_report["obstacle_ids"] = set(parse_point_ids(raw))
            active_report["pending_ids"] = set()
            st.rerun()
    elif init_mode == "上传掩膜图（黑障碍/白可达）":
        mask = st.file_uploader(
            "上传掩膜图（png/jpg，建议比例 28:30）",
            type=["png", "jpg", "jpeg"],
            key=f"mask_{active_report_key}",
        )
        if mask is not None and st.button("应用掩膜", key=f"apply_mask_{active_report_key}", type="primary"):
            try:
                active_report["obstacle_ids"] = obstacle_ids_from_mask(mask.getvalue())
                active_report["pending_ids"] = set()
                st.rerun()
            except Exception as exc:
                st.error(str(exc))
    render_stage_subpanel_end()
    render_section_subsection(
        "地图修正",
        "这一层处理点选、框选、颜色预览以及右侧状态写入，先确认选区，再决定具体写入什么状态。",
        "2.2",
    )
    render_stage_subpanel_start("地图修正", "2.2")
    render_switcher_panel(
        "地图日期切换",
        "map_report_switch",
        report_options,
        st.session_state.active_report_key,
    )
    active_report_key = st.session_state.active_report_key
    active_report = report_store[active_report_key]
    parsed = sync_report_pollution_state(active_report)
    polluted_ids = get_report_polluted_ids(active_report)

    render_note_bar("支持多次拖拽累积选中；颜色先预览，点击按钮后再确认写入状态。", soft=True)
    render_terminology_legend()
    map_col, action_col = st.columns([1.42, 1.24], gap="medium")
    with map_col:
        fig = build_plotly_grid(
            polluted_ids,
            set(active_report["obstacle_ids"]),
            set(active_report["pending_ids"]),
            set(active_report["selected_ids"]),
        )
        selection_event = st.plotly_chart(
            fig,
            use_container_width=True,
            on_select="rerun",
            selection_mode=["box", "lasso"],
            config={
                "responsive": True,
                "scrollZoom": False,
                "displaylogo": False,
                "doubleClick": False,
            },
            key=f"grid_select_{active_report_key}",
        )
    new_selected_ids = extract_selected_ids(selection_event)
    if has_selection_payload(selection_event) and new_selected_ids:
        active_report["selected_ids"] = merge_selected_ids(
            set(active_report["selected_ids"]),
            new_selected_ids,
        )
    selected_ids = set(active_report["selected_ids"])

    with action_col:
        st.markdown(
            """
            <div class="toolbar-side-shell">
                <div class="toolbar-side-title">2.3 状态写入</div>
                <div class="toolbar-side-note">右侧工具栏用于把当前选区写入污染、无污染、障碍或待测状态。</div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown(
            f'<div class="toolbar-side-count">当前选区：{len(selected_ids)} 个点位</div>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="toolbar-row-title">2.3.0 当前选中 ID</div>', unsafe_allow_html=True)
        render_toolbar_selected_ids_panel(sorted(selected_ids))

        st.markdown('<div class="toolbar-row-title">2.3.1 污染判定</div>', unsafe_allow_html=True)
        row1 = st.columns(2, gap="small")
        with row1[0]:
            if st.button(
                "设为污染",
                key=f"set_polluted_{active_report_key}",
                disabled=not selected_ids,
                help="将当前选中的格子确认写入污染点位。",
                type="primary",
                use_container_width=True,
            ):
                active_report["polluted_ids"] = apply_pollution_action(
                    "set_polluted",
                    selected_ids,
                    polluted_ids,
                )
                active_report["obstacle_ids"] = set(active_report["obstacle_ids"]) - selected_ids
                active_report["pending_ids"] = set(active_report["pending_ids"]) - selected_ids
                active_report["selected_ids"] = set()
                st.rerun()
        with row1[1]:
            if st.button(
                "设为无污染",
                key=f"set_clean_{active_report_key}",
                disabled=not selected_ids,
                help="将当前选中的格子从污染集合中移出，视为无污染可达区域。",
                type="primary",
                use_container_width=True,
            ):
                active_report["polluted_ids"] = apply_pollution_action(
                    "set_clean",
                    selected_ids,
                    polluted_ids,
                )
                active_report["obstacle_ids"] = set(active_report["obstacle_ids"]) - selected_ids
                active_report["pending_ids"] = set(active_report["pending_ids"]) - selected_ids
                active_report["selected_ids"] = set()
                st.rerun()

        st.markdown('<div class="toolbar-row-spacer"></div>', unsafe_allow_html=True)

        st.markdown('<div class="toolbar-row-title">2.3.2 区域定义</div>', unsafe_allow_html=True)
        row2 = st.columns(3, gap="small")
        with row2[0]:
            if st.button(
                "设为障碍",
                key=f"set_obstacle_{active_report_key}",
                disabled=not selected_ids,
                help="将当前选中格子设置为障碍物，并从污染/待测状态中移除。",
                type="primary",
                use_container_width=True,
            ):
                obstacle_ids, pending_ids = apply_selection_action(
                    "set_obstacle",
                    selected_ids,
                    set(active_report["obstacle_ids"]),
                    set(active_report["pending_ids"]),
                )
                active_report["obstacle_ids"] = obstacle_ids
                active_report["pending_ids"] = pending_ids
                active_report["polluted_ids"] = polluted_ids - selected_ids
                active_report["selected_ids"] = set()
                st.rerun()
        with row2[1]:
            if st.button(
                "设为待测",
                key=f"set_pending_{active_report_key}",
                disabled=not selected_ids,
                help="将当前选中格子标记为待测区域，后续再复核污染状态。",
                type="primary",
                use_container_width=True,
            ):
                obstacle_ids, pending_ids = apply_selection_action(
                    "set_pending",
                    selected_ids,
                    set(active_report["obstacle_ids"]),
                    set(active_report["pending_ids"]),
                )
                active_report["obstacle_ids"] = obstacle_ids
                active_report["pending_ids"] = pending_ids
                active_report["polluted_ids"] = polluted_ids - selected_ids
                active_report["selected_ids"] = set()
                st.rerun()
        with row2[2]:
            if st.button(
                "清除定义",
                key=f"clear_selected_{active_report_key}",
                disabled=not selected_ids,
                help="清空当前选中格子的障碍/待测标记，但不会自动改回污染状态。",
                type="secondary",
                use_container_width=True,
            ):
                obstacle_ids, pending_ids = apply_selection_action(
                    "clear",
                    selected_ids,
                    set(active_report["obstacle_ids"]),
                    set(active_report["pending_ids"]),
                )
                active_report["obstacle_ids"] = obstacle_ids
                active_report["pending_ids"] = pending_ids
                active_report["selected_ids"] = set()
                st.rerun()

        st.markdown('<div class="toolbar-row-title">2.3.3 快速操作</div>', unsafe_allow_html=True)
        row3 = st.columns([1.0, 1.0, 1.0], gap="small")
        with row3[0]:
            if st.button(
                "取消选中",
                key=f"clear_staged_{active_report_key}",
                disabled=not selected_ids,
                help="只清空当前暂存选区，不改动污染/障碍/待测结果。",
                type="secondary",
                use_container_width=True,
            ):
                active_report["selected_ids"] = set()
                st.rerun()
        with row3[1]:
            if st.button(
                "默认障碍",
                key=f"restore_default_{active_report_key}",
                help="将障碍物分布恢复到系统默认布局。",
                type="secondary",
                use_container_width=True,
            ):
                active_report["obstacle_ids"] = default_obstacle_ids()
                active_report["pending_ids"] = set()
                st.rerun()
        with row3[2]:
            if st.button(
                "重置污染",
                key=f"reset_pollution_{active_report_key}",
                help="按当前报告文本重新覆盖污染点位，丢弃人工修正后的污染结果。",
                type="secondary",
                use_container_width=True,
            ):
                active_report["polluted_ids"] = set(parsed.floor_ids)
                active_report["selected_ids"] = set()
                active_report["pollution_source_signature"] = build_text_signature(
                    str(active_report["report_text"])
                )
                st.rerun()
        st.markdown('<div class="toolbar-side-footer"></div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    render_stage_panel_end()
    render_major_section_divider()

    selected_list = sorted(selected_ids)
    accessible_clean_count = TOTAL - len(
        polluted_ids | set(active_report["obstacle_ids"]) | set(active_report["pending_ids"])
    )
    render_stage_panel_start(
        "日期结果核查",
        "三、日期结果",
        "展示当前激活日期的识别结果、人工修正后的统计以及污染点位明细，便于逐日报告核对。",
        chips=[
            (
                "当前日期",
                str(active_report.get("date") or "未识别日期"),
                "结果面板始终对应当前激活的那一期报告",
            ),
            (
                "当前报告",
                str(active_report.get("file_name", "-")),
                "报告文本修改后，这里的统计会同步更新",
            ),
            (
                "当前污染点位",
                str(len(polluted_ids)),
                "综合报告识别和人工修正后的最终结果",
            ),
        ],
    )
    render_daily_result_section_map()
    render_stage_meta_strip(
        [
            ("超标手套孔", str(len(parsed.glove_names))),
            ("障碍物点位", str(len(active_report["obstacle_ids"]))),
            ("待测区域", str(len(active_report["pending_ids"]))),
            ("当前无污染点位", str(accessible_clean_count)),
        ]
    )
    st.markdown('<div class="result-section-offset"></div>', unsafe_allow_html=True)
    st.markdown('<div class="result-title">日期结果</div>', unsafe_allow_html=True)
    render_result_hero(len(polluted_ids), len(parsed.glove_names), len(selected_list))
    render_analytics_subsection(
        "统计概览",
        "这里先给出当前日期下的核心数量结果，方便快速判断污染规模和可达区域状态。",
        "3.1",
    )
    top_metrics = st.columns(4)
    with top_metrics[0]:
        render_metric_card("当前污染点位", len(polluted_ids), alert=True)
    with top_metrics[1]:
        render_metric_card("超标手套孔", len(parsed.glove_names), alert=True)
    with top_metrics[2]:
        render_metric_card("障碍物点位", len(active_report["obstacle_ids"]))
    with top_metrics[3]:
        render_metric_card("待测区域点位", len(active_report["pending_ids"]))
    st.markdown('<div class="section-gap-sm"></div>', unsafe_allow_html=True)
    render_analytics_subsection(
        "结果明细",
        "这一层把识别和修正后的结果展开成可核对明细，包括污染点位、手套孔和行列映射。",
        "3.2",
    )
    map_col, detail_col = st.columns([1.42, 0.98], gap="large")
    with map_col:
        st.plotly_chart(
            build_result_snapshot_grid(
                polluted_ids,
                set(active_report["obstacle_ids"]),
                set(active_report["pending_ids"]),
            ),
            use_container_width=True,
            config={
                "responsive": True,
                "scrollZoom": False,
                "displaylogo": False,
                "doubleClick": False,
            },
        )
    with detail_col:
        render_detection_details_panel(
            sorted(polluted_ids),
            parsed.glove_names,
            [
                ("当前报告", str(active_report.get("file_name", "-"))),
                ("识别日期", str(active_report.get("date") or "未识别日期")),
                ("无污染点位", str(accessible_clean_count)),
            ],
        )
        if active_report.get("load_error"):
            st.error(f"文件解析失败：{active_report.get('load_error')}")
    render_stage_panel_end()
    render_major_section_divider()

    summary_rows = build_cross_report_summary(report_options, report_store)
    if summary_rows:
        latest_summary = summary_rows[-1]
        peak_pollution = max(int(row["污染地面点位数"]) for row in summary_rows)
        total_added = sum(int(row["新增点位数"]) for row in summary_rows)
        total_removed = sum(int(row["减少点位数"]) for row in summary_rows)
        summary_chips = [
            (
                "汇总期数",
                str(len(summary_rows)),
                "按日期顺序串联多份报告，形成横向变化对照表",
            ),
            (
                "最新污染点位",
                str(latest_summary["污染地面点位数"]),
                f"对应日期 {latest_summary['日期']}，可与地图和预测区联动查看",
            ),
            (
                "历史峰值",
                str(peak_pollution),
                "用于快速识别污染高位区间，辅助判断阶段性扩张",
            ),
        ]
    else:
        summary_chips = [
            ("汇总期数", "0", "当前还没有可用于横向汇总的历史报告"),
            ("最新污染点位", "-", "上传报告后会在这里显示最新一期统计"),
            ("历史峰值", "-", "随着报告积累自动更新"),
        ]

    render_stage_panel_start(
        "日期汇总对照",
        "四、普查结果按日期汇总",
        "按日期串联所有报告，形成一张普查结果汇总表，直接查看各期污染点位、手套孔数量和前后变化。",
        chips=summary_chips,
    )
    render_section_path_map(
        [
            ("4.1", "汇总概览", "先看期数、最新结果、历史峰值以及累计增减。"),
            ("4.2", "日期对照表", "再用一张表横向比较每个日期下的普查结果。"),
        ]
    )
    if summary_rows:
        render_section_subsection(
            "汇总概览",
            "这一层先看期数、最新结果、历史峰值以及累计新增减少，快速判断整体变化趋势。",
            "4.1",
        )
        render_stage_meta_strip(
            [
                ("最新日期", str(latest_summary["日期"])),
                ("累计新增点位", str(total_added)),
                ("累计减少点位", str(total_removed)),
                ("当前变化", str(latest_summary["较前一日变化"])),
            ]
        )
        render_section_subsection(
            "日期对照表",
            "这一层用表格串联所有日期结果，便于快速比对每一期的污染数量和新增减少变化。",
            "4.2",
        )
        render_stage_subpanel_start("跨日期汇总表与导出", "4.2")
        render_export_button("导出多日期汇总 CSV", summary_rows, "cross_report_summary.csv", "export_summary_rows")
        st.dataframe(summary_rows, use_container_width=True, hide_index=True)
        render_note_bar("按日期排序；“较前一日变化”以当前行与上一行污染地面点位数对比。", soft=True)
        render_stage_subpanel_end()
    else:
        render_section_subsection(
            "汇总概览",
            "这一层会在有多期报告后展示整体汇总状态，目前还没有足够的数据。",
            "4.1",
        )
        render_section_subsection(
            "日期对照表",
            "这一层会在有多期报告后自动生成横向对照表，目前还没有足够的数据。",
            "4.2",
        )
        render_stage_subpanel_start("跨日期汇总表与导出", "4.2")
        render_note_bar("暂无可汇总的数据。", soft=True)
        render_stage_subpanel_end()
    render_stage_panel_end()
    render_major_section_divider()

    render_analytics_section(report_options, report_store)


if __name__ == "__main__":
    main()
