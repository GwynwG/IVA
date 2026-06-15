from docx import Document

template_path = r"C:\Users\罗兆东\Desktop\4XTS\AI高价值场景应用\污染普查数据可视化\示例文件\20250701.docx"
doc = Document(template_path)

print("=== 段落 ===")
for i, para in enumerate(doc.paragraphs):
    print(f"段落 {i}: {para.text[:100]}")

print("\n=== 表格 ===")
for t, table in enumerate(doc.tables):
    print(f"表格 {t}: {len(table.rows)} 行, {len(table.columns)} 列")
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  表格[{t}] 行[{r}] 列[{c}]: {text}")

print("\n=== 全文文本 ===")
full_text = []
for para in doc.paragraphs:
    full_text.append(para.text)
print("\n".join(full_text))